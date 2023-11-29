
import streamlit as st
from openai import OpenAI
import uuid
from PIL import Image
import numpy as np
import pdfplumber
import pandas as pd
from docx import Document
import base64
import re
import os
import difflib


# ã‚µãƒ¼ãƒ“ã‚¹åã‚’è¡¨ç¤ºã™ã‚‹
st.sidebar.title("AI Assistant")

# åˆå›ãƒ­ã‚°ã‚¤ãƒ³èªè¨¼
if "authenticated" not in st.session_state:
    st.session_state["authenticated"] = False

if not st.session_state["authenticated"]:
    user_id = st.text_input("ãƒ¦ãƒ¼ã‚¶ãƒ¼IDã‚’å…¥åŠ›ã—ã¦ãã ã•ã„:")
    password = st.text_input("ãƒ‘ã‚¹ãƒ¯ãƒ¼ãƒ‰ã‚’å…¥åŠ›ã—ã¦ãã ã•ã„:", type="password")
    if st.button("ãƒ­ã‚°ã‚¤ãƒ³"):
        if user_id == "admin" and password == "LLM@2023":
            st.session_state["authenticated"] = True
            st.success("ãƒ­ã‚°ã‚¤ãƒ³æˆåŠŸ!")
            if st.button("ç¶šã‘ã‚‹"):
                pass
        else:
            st.error("èª¤ã£ãŸãƒ¦ãƒ¼ã‚¶ãƒ¼IDã¾ãŸã¯ãƒ‘ã‚¹ãƒ¯ãƒ¼ãƒ‰ã§ã™ã€‚")

if st.session_state["authenticated"]:
    # Create a unique key for the widget
    unique_key = str(uuid.uuid4())

    # Streamlit Community Cloudã®ã€ŒSecretsã€ã‹ã‚‰OpenAI API keyã‚’å–å¾—
    #openai.api_key = st.secrets.OpenAIAPI.openai_api_key
    api_key = st.secrets.OpenAIAPI.openai_api_key

    # st.session_stateã‚’ä½¿ã„ãƒ¡ãƒƒã‚»ãƒ¼ã‚¸ã®ã‚„ã‚Šã¨ã‚Šã‚’ä¿å­˜
    if "messages" not in st.session_state:
        st.session_state["messages"] = [
            {"role": "system", "content": "You are the best AI assistant in the world."}
        ]

    # ApiClientã‚¤ãƒ³ã‚¹ã‚¿ãƒ³ã‚¹ã‚’ä½œæˆ
    client = OpenAI(api_key = api_key)

    if "user_input" not in st.session_state:
        st.session_state["user_input"] = ""

#â†“ãƒˆãƒ¼ã‚¯ãƒ³æ•°ä¸Šé™ç›£è¦–ã®ãŸã‚ã®ã‚³ãƒ¼ãƒ‰
    # ãƒˆãƒ¼ã‚¯ãƒ³ä½¿ç”¨é‡ã‚’è¨˜éŒ²ã™ã‚‹ã‚»ãƒƒã‚·ãƒ§ãƒ³ã‚¹ãƒ†ãƒ¼ãƒˆã®åˆæœŸåŒ–
    #if 'token_usage' not in st.session_state:
    #    st.session_state['token_usage'] = 0

    # ãƒˆãƒ¼ã‚¯ãƒ³ä½¿ç”¨é‡ã®ä¸Šé™
    #TOKEN_LIMIT = 100000  # ä¾‹ã¨ã—ã¦100,000ãƒˆãƒ¼ã‚¯ãƒ³ã‚’ä¸Šé™ã¨ã—ã¾ã™

    # ãƒˆãƒ¼ã‚¯ãƒ³ä½¿ç”¨é‡ã‚’æ›´æ–°ã™ã‚‹é–¢æ•°
    #def update_token_usage(new_tokens):
    #    st.session_state['token_usage'] += new_tokens
    #    # ä¸Šé™ã«é”ã—ãŸå ´åˆã®è­¦å‘Šè¡¨ç¤º
    #    if st.session_state['token_usage'] >= TOKEN_LIMIT:
    #        st.warning('ãƒˆãƒ¼ã‚¯ãƒ³ä½¿ç”¨é‡ã®ä¸Šé™ã«é”ã—ã¾ã—ãŸã€‚')

    # APIãƒªã‚¯ã‚¨ã‚¹ãƒˆã‚’è¡Œã†é–¢æ•°ï¼ˆä»®ï¼‰
    #def make_api_request():
    #    # ã“ã“ã§APIãƒªã‚¯ã‚¨ã‚¹ãƒˆã‚’è¡Œã„ã€ãƒ¬ã‚¹ãƒãƒ³ã‚¹ã‚’å–å¾—ã—ã¾ã™
    #    # ä»®ã«ãƒˆãƒ¼ã‚¯ãƒ³æ•°ã‚’50ã¨ã—ã¾ã™
    #    used_tokens = 50
    #    return used_tokens

    # ãƒœã‚¿ãƒ³ã§APIãƒªã‚¯ã‚¨ã‚¹ãƒˆã‚’ãƒˆãƒªã‚¬ãƒ¼ã™ã‚‹
    #if st.button('APIãƒªã‚¯ã‚¨ã‚¹ãƒˆå®Ÿè¡Œ'):
    #    tokens = make_api_request()
    #    update_token_usage(tokens)

    # ç¾åœ¨ã®ãƒˆãƒ¼ã‚¯ãƒ³ä½¿ç”¨é‡ã‚’è¡¨ç¤º
    #st.write(f'ç¾åœ¨ã®ãƒˆãƒ¼ã‚¯ãƒ³ä½¿ç”¨é‡: {st.session_state["token_usage"]}')
#â†‘ãƒˆãƒ¼ã‚¯ãƒ³æ•°ä¸Šé™ç›£è¦–ã®ãŸã‚ã®ã‚³ãƒ¼ãƒ‰


    #ãƒˆãƒ¼ã‚¯ãƒ³æ•°ã‚«ã‚¦ãƒ³ãƒˆï¼ˆä¿®æ­£å‰ï¼‰
    #def count_tokens(text):
    #    response = client.completions.create(model="text-davinci-002", messages=[{"role": "system", "content": text}])
    #    token_count = response['usage']['total_tokens']
    #    return token_count


    #ãƒˆãƒ¼ã‚¯ãƒ³æ•°ã‚«ã‚¦ãƒ³ãƒˆï¼ˆä¿®æ­£ä¸­ï¼‰
    #def count_tokens(text):
    #    response = client.completions.create({
    #        "model": "text-davinci-002",  # or your preferred model
    #        "prompt": user_input,        # the text you want to process
    #        "max_tokens": 1              # maximum number of tokens to generate
    #    })
    #    #ãƒ¬ã‚¹ãƒãƒ³ã‚¹ã®å†…å®¹ã‚’è¡¨ç¤º
    #    st.write(response)

        # 'Completion' ã‚ªãƒ–ã‚¸ã‚§ã‚¯ãƒˆã‹ã‚‰å¿…è¦ãªæƒ…å ±ã‚’å–å¾—
        #token_count = response.choices[0].usage.total_tokens
        #return token_count

    def get_binary_file_downloader_html(bin_file, file_label="File"):
        with open(bin_file, "rb") as f:
            data = f.read()
        bin_str = base64.b64encode(data).decode()
        href = f'<a href="data:application/octet-stream;base64,{bin_str}" download="{file_label}.docx">Download {file_label}</a>'
        return href

    # ï¼ˆã‚¢ãƒƒãƒ—ãƒ‡ãƒ¼ãƒˆå‰ã®ã‚³ãƒ¼ãƒ‰ï¼‰ãƒãƒ£ãƒƒãƒˆãƒœãƒƒãƒˆã¨ã‚„ã‚Šã¨ã‚Šã™ã‚‹é–¢æ•°
    #def communicate(user_input, bot_response_placeholder, model, temperature, top_p):
    #    messages = st.session_state["messages"]
    #    user_message = {"role": "user", "content": user_input}
    #    messages.append(user_message)

    #    # Temporary variable to store chunks
    #    complete_response = ""

    #    # Get the response from ChatCompletion in streaming mode
    #    for chunk in client.completions.create(
    #        model=model,
    #        messages=messages,
    #        temperature=temperature,
    #        top_p=top_p,
    #        stream=True
    #    ):
    #        content = chunk["choices"][0].get("delta", {}).get("content")
    #        if content is not None:
    #            # Accumulate content and update the bot's response in real time
    #            complete_response += content
    #            formatted_response = complete_response.replace("\n", "<br>")
    #            indented_response = "".join([f"<div style='margin-left: 20px; white-space: pre-wrap;'>{line}</div>" for line in complete_response.split('\n')]) # ã‚¤ãƒ³ãƒ‡ãƒ³ãƒˆã§å›ç­”
    #            bot_response_placeholder.markdown(indented_response, unsafe_allow_html=True)

    #    # After all chunks are received, add the complete response to the chat history
    #    if complete_response:
    #        bot_message = {"role": "assistant", "content": complete_response}
    #        messages.append(bot_message)

    #   # Reset the messages after the chat
    #    messages = [{"role": "system", "content": "You are the best AI assistant in the world."}]

    #    return complete_response


# ï¼ˆä½œæ¥­ä¸­ï¼‰ãƒãƒ£ãƒƒãƒˆãƒœãƒƒãƒˆã¨ã‚„ã‚Šã¨ã‚Šã™ã‚‹é–¢æ•°ï¼ˆã‚¹ãƒˆãƒªãƒ¼ãƒ ãƒ¢ãƒ¼ãƒ‰ï¼‰
# def communicate(user_input, bot_response_placeholder, model, temperature, top_p):
#     messages = st.session_state["messages"]
#     user_message = {"role": "user", "content": user_input}
#     messages.append(user_message)

#     complete_response = ""

#     # Get the response from ChatCompletion
#     response = client.chat.completions.create(
#         model=model,
#         messages=messages,
#         temperature=temperature,
#         max_tokens=4000,
#         top_p=top_p,
#         stream=True
#     )
#     for chunk in response:
#         # 'choices'å±æ€§ã‹ã‚‰ãƒ¬ã‚¹ãƒãƒ³ã‚¹ãƒ†ã‚­ã‚¹ãƒˆã‚’å–å¾—
#         if chunk.choices:
#             response_text = chunk.choices[0].text
#             if response_text is not None:
#                 # Accumulate content and update the bot's response in real time
#                 complete_response += response_text
#                 formatted_response = complete_response.replace("\n", "<br>")
#                 indented_response = "".join([f"<div style='margin-left: 20px; white-space: pre-wrap;'>{line}</div>" for line in complete_response.split('\n')]) # ã‚¤ãƒ³ãƒ‡ãƒ³ãƒˆã§å›ç­”
#                 bot_response_placeholder.markdown(indented_response, unsafe_allow_html=True)

#     # After all chunks are received, add the complete response to the chat history
#     if complete_response:
#         bot_message = {"role": "assistant", "content": complete_response}
#         messages.append(bot_message)

#     # Reset the messages after the chat
#     messages = [{"role": "system", "content": "You are the best AI assistant in the world."}]

#     return complete_response


    # ãƒãƒ£ãƒƒãƒˆãƒœãƒƒãƒˆã¨ã‚„ã‚Šã¨ã‚Šã™ã‚‹é–¢æ•°
    def communicate(user_input, bot_response_placeholder, model, temperature, top_p):
        messages = st.session_state["messages"]
        user_message = {"role": "user", "content": user_input}
        messages.append(user_message)

        complete_response = ""

        # Get the response from ChatCompletion
        response = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=temperature,
            max_tokens=2000,
            top_p=top_p,
        )
        response_message = response.choices[0].message.content
        if response_message is not None:
            # Accumulate content and update the bot's response in real time
            complete_response += response_message
            formatted_response = complete_response.replace("\n", "<br>")
            indented_response = "".join([f"<div style='margin-left: 20px; white-space: pre-wrap;'>{line}</div>" for line in complete_response.split('\n')]) # ã‚¤ãƒ³ãƒ‡ãƒ³ãƒˆã§å›ç­”
            bot_response_placeholder.markdown(indented_response, unsafe_allow_html=True)

        # After all chunks are received, add the complete response to the chat history
        if complete_response:
            bot_message = {"role": "assistant", "content": complete_response}
            messages.append(bot_message)

        # Reset the messages after the chat
        messages = [{"role": "system", "content": "You are the best AI assistant in the world."}]

        return complete_response

#    def process_response(generated_text, user_input):
#        # åˆ†å‰²ã—ã¦ã€ä¿®æ­£å¾Œã®å…¨æ–‡ã¨ä¿®æ­£ç®‡æ‰€ãƒªã‚¹ãƒˆã‚’æŠ½å‡º
#        response_lines = generated_text.split("\n")
#        corrected_full_text = response_lines[0]  # æœ€åˆã®è¡Œã‚’ã€Œä¿®æ­£å¾Œå…¨æ–‡ã€ã¨ä»®å®š
#        correction_list = response_lines[1:]  # æ®‹ã‚Šã®è¡Œã‚’ã€Œä¿®æ­£ç®‡æ‰€ãƒªã‚¹ãƒˆã€ã¨ä»®å®š

#        # å…ƒã®ãƒ†ã‚­ã‚¹ãƒˆã¨ä¿®æ­£å¾Œã®ãƒ†ã‚­ã‚¹ãƒˆã®æ¯”è¼ƒ
#        corrected_full_text_words = corrected_full_text.split()
#        user_input_words = user_input.split()
#        bolded_text = ""

        # å„å˜èªã‚’æ¯”è¼ƒã—ã¦ã€å¤‰æ›´ã•ã‚ŒãŸéƒ¨åˆ†ã‚’å¼·èª¿
#        for word in corrected_full_text_words:
#            if word in user_input_words:
#                bolded_text += word + " "
#            else:
#                bolded_text += "**" + word + "** "
#
#        return bolded_text, correction_list



    # ã‚µã‚¤ãƒ‰ãƒãƒ¼ã§æ©Ÿèƒ½ã‚’é¸æŠ
    selected_option = st.sidebar.selectbox(
        "æ©Ÿèƒ½ã‚’é¸æŠã—ã¦ãã ã•ã„",
        ["é¸æŠã—ã¦ãã ã•ã„", "Q&A", "Translation", "Proofreading", "Excel Formula Analysis", "VBA Analysis", "Data Analysis"],
        index=0, # ãƒ‡ãƒ•ã‚©ãƒ«ãƒˆå€¤ã¨ã—ã¦ã€Œé¸æŠã—ã¦ãã ã•ã„ã€ã‚’è¨­å®š
        key="selectbox_key"  # å›ºå®šã®ã‚­ãƒ¼ã‚’æŒ‡å®šã™ã‚‹
    )

    # ã‚¿ã‚¤ãƒˆãƒ«ã€Œã‚ªãƒ—ã‚·ãƒ§ãƒ³ã€ã‚’è¿½åŠ 
    st.sidebar.header("ã‚ªãƒ—ã‚·ãƒ§ãƒ³")

    # ãƒ¢ãƒ‡ãƒ«ã®é¸æŠã¨ãã®è£œè¶³æƒ…å ±
    with st.sidebar.expander("ãƒ¢ãƒ‡ãƒ«  ğŸ›ˆ"):
        st.write(
        """gpt-4ï¼ˆæ¨å¥¨ï¼‰ã¯ã€é«˜å“è³ªãªå›ç­”ã‚’å‡ºåŠ›ã—ã¾ã™ã€‚å…¥åŠ›ãƒ»å‡ºåŠ›ã®åˆè¨ˆã§ç´„8,000ãƒˆãƒ¼ã‚¯ãƒ³ã¾ã§å‡¦ç†å¯èƒ½ã§ã™ã€‚gpt-3.5-turbo-16kã¯ã€gpt-4ã¨æ¯”è¼ƒã™ã‚‹ã¨å›ç­”ã®è³ªã¯ä¸‹ãŒã‚Šã¾ã™ãŒã€å…¥åŠ›ãƒ»å‡ºåŠ›ã®åˆè¨ˆã§ç´„16,000ãƒˆãƒ¼ã‚¯ãƒ³ã¾ã§å‡¦ç†ã§ãã€gpt-4ã«æ¯”ã¹é«˜é€Ÿã§å›ç­”ã®å‡ºåŠ›ãŒå¯èƒ½ã§ã™ã€‚
        """)
        model = st.selectbox(
        "ãƒ¢ãƒ‡ãƒ«ã‚’é¸æŠã—ã¦ãã ã•ã„",
        ["gpt-4", "gpt-3.5-turbo-16k"],
        key="model_selectbox_key"  # å›ºå®šã®ã‚­ãƒ¼ã‚’æŒ‡å®šã™ã‚‹
    )

    # Temperatureã‚¹ãƒ©ã‚¤ãƒ€ãƒ¼ã¨ãã®è£œè¶³æƒ…å ±
    with st.sidebar.expander("Temperature  ğŸ›ˆ"):
        st.write("Temperatureï¼ˆæ¸©åº¦ï¼‰:ãƒ¢ãƒ‡ãƒ«ã®å‡ºåŠ›ã®ã€Œç¢ºä¿¡åº¦ã€ã¾ãŸã¯ã€Œå¤šæ§˜æ€§ã€ã‚’åˆ¶å¾¡ã—ã¾ã™ã€‚å€¤ãŒé«˜ã„ã¨ãƒ¢ãƒ‡ãƒ«ã®å‡ºåŠ›ã¯å¤šæ§˜æ€§ãŒå¢—ã—ã€äºˆæ¸¬ã¯ã‚ˆã‚Šãƒ©ãƒ³ãƒ€ãƒ ã«ãªã‚Šã¾ã™ã€‚é€†ã«ã€å€¤ãŒä½ã„ã¨ãƒ¢ãƒ‡ãƒ«ã®å‡ºåŠ›ã¯ã‚ˆã‚Šç¢ºä¿¡åº¦ãŒé«˜ããªã‚Šã€æœ€ã‚‚ç¢ºç‡çš„ã«é«˜ã„çµæœã‚’é¸ã³ã‚„ã™ããªã‚Šã¾ã™ã€‚ã€æ¨å¥¨å€¤:0.10ã€‘")
        temperature = st.slider("", 0.0, 2.0, 0.1, 0.01)

    # Top_Pã‚¹ãƒ©ã‚¤ãƒ€ãƒ¼ã¨ãã®è£œè¶³æƒ…å ±
    with st.sidebar.expander("Top_P  ğŸ›ˆ"):
        st.write("Top_P: æ¸©åº¦ã¨åŒæ§˜ã«ã€ã“ã‚Œã¯ãƒ©ãƒ³ãƒ€ãƒ æ€§ã‚’åˆ¶å¾¡ã—ã¾ã™ãŒã€åˆ¥ã®æ–¹æ³•ã‚’ä½¿ç”¨ã—ã¾ã™ã€‚Top_P ã‚’ä¸‹ã’ã‚‹ã¨ã€ã‚ˆã‚Šå¯èƒ½æ€§ãŒé«˜ã„å›ç­”ã«çµã‚Šè¾¼ã¾ã‚Œã¾ã™ã€‚Top_P ã‚’ä¸Šã’ã‚‹ã¨ã€ç¢ºç‡ãŒé«˜ã„å›ç­”ã¨ä½ã„å›ç­”ã®ä¸¡æ–¹ã‹ã‚‰é¸æŠã•ã‚Œã‚‹ã‚ˆã†ã«ãªã‚Šã¾ã™ã€‚ã€æ¨å¥¨å€¤:0.50ã€‘")
        top_p = st.slider("", 0.0, 1.0, 0.5, 0.01)

    # ç´¯ç©ãƒˆãƒ¼ã‚¯ãƒ³æ•°ãƒªã‚»ãƒƒãƒˆãƒœã‚¿ãƒ³ã®è¨­ç½®
    #if st.sidebar.button("ãƒˆãƒ¼ã‚¯ãƒ³æ•°ãƒªã‚»ãƒƒãƒˆ"):
    #    st.session_state["messages"] = [
    #        {"role": "system", "content": "You are the best AI assistant in the world."}
    #    ]

    # ã€ŒãŠå•ã„åˆã‚ã›ã€ãƒã‚¤ãƒ‘ãƒ¼ãƒªãƒ³ã‚¯ã®è¨­ç½®
    def create_mailto_link():
        to_address = "kazuki.takahashi@front-ia.com,katakahashi@pictet.com"
        cc_address = "hiroyuki.tsuchida@front-ia.com"
        subject = "AI Assistant"
        return f"mailto:{to_address}?subject={subject}&cc={cc_address}"

    mailto_link = create_mailto_link()
    st.sidebar.markdown(f'<a href="{mailto_link}" target="_blank">ãŠå•ã„åˆã‚ã›</a>', unsafe_allow_html=True)

    # (æº–å‚™ä¸­)ãƒ¦ãƒ¼ã‚¶ãƒ¼ã‚¢ãƒ³ã‚±ãƒ¼ãƒˆ
    #st.sidebar.markdown("""
    #[ãŠå•ã„åˆã‚ã›](https://ai-assistant-inquiries-8sft4gmafubshjqsrzx6m2.streamlit.app/)
    #""")

    # ãƒãƒ¼ã‚¸ãƒ§ãƒ³æƒ…å ±è¡¨ç¤ºï¼ˆãƒªãƒªãƒ¼ã‚¹ãƒãƒ¼ãƒˆã¸ã®ãƒã‚¤ãƒ‘ãƒ¼ãƒªãƒ³ã‚¯ï¼‰
    st.sidebar.markdown("""
    [v2.0.0](https://ai-assistant-releasenote-mfjkhzwcdpy9p33km6tffg.streamlit.app/)
    """)

   #Wordãƒ•ã‚¡ã‚¤ãƒ«å½¢å¼ã§ã®å‡ºåŠ›ã‚’å®šç¾©
    def create_word_doc(text):
        doc = Document()
        doc.add_paragraph(text)
        output_path = "/tmp/translated_text.docx"
        doc.save(output_path)
        return output_path

    #Wordãƒ•ã‚¡ã‚¤ãƒ«å½¢å¼ã§ã®å‡ºåŠ›ã‚’å®šç¾©ï¼ˆãƒ•ã‚¡ã‚¤ãƒ«ã‚¿ã‚¤ãƒˆãƒ«ã‚’åŸæ–‡ã‚¿ã‚¤ãƒˆãƒ«ã‹ã‚‰å–å¾—ã™ã‚‹ã‚ˆã†ä¿®æ­£ä¸­ï¼‰
    #def create_word_doc(text):
        # æœ€åˆã®è¡Œã‚’å–å¾—
#        first_line = text.split("\n")[0].strip()
#        print(f"Debug: First Line = {first_line}")

        # æ­£è¦è¡¨ç¾ã‚’ä½¿ç”¨ã—ã¦ã‚¿ã‚¤ãƒˆãƒ«éƒ¨åˆ†ã‚’å–å¾—
#        match = re.search(r'\d{2}:\d{2} (.+)', first_line)
#        if match:
#            title = match.group(1)
#        else:
#            title = first_line

#        print(f"Debug: Title = {title}")


        # ãƒ•ã‚¡ã‚¤ãƒ«åã¨ã—ã¦ä½¿ç”¨ã§ããªã„æ–‡å­—ã‚’å–ã‚Šé™¤ã
#        valid_filename = re.sub(r"[^a-zA-Z0-9]", "_", title)
#        valid_filename = re.sub(r"_+", "_", valid_filename)

        # ä¸€å®šã®é•·ã•ã«åˆ¶é™ã™ã‚‹ (ä¾‹: 20æ–‡å­—)
#        valid_filename = valid_filename[:20] + ".docx"

#        print(f"Debug: Valid Filename = {valid_filename}")

#        doc = Document()
#        doc.add_paragraph(text)
#        output_path = f"/tmp/{valid_filename}"
#        doc.save(output_path)
#        return output_path


    # æ©Ÿèƒ½ã«å¿œã˜ãŸUIã®è¡¨ç¤º
    if selected_option == "é¸æŠã—ã¦ãã ã•ã„":
        pass  # ä½•ã‚‚è¡¨ç¤ºã—ãªã„

    elif selected_option == "Q&A":
        # Build the user interface
        st.title("Q&A")

        # ç•™æ„ç‚¹ã®è¡¨ç¤º
        st.markdown('<span style="color:red">***å€‹äººæƒ…å ±ã‚„æ©Ÿå¯†æƒ…å ±ã¯å…¥åŠ›ã—ãªã„ã§ãã ã•ã„**</span>', unsafe_allow_html=True)

        # ãƒ¦ãƒ¼ã‚¶ãƒ¼å…¥åŠ›ã‚’åˆæœŸåŒ–
        user_input = ""
        uploaded_file = ""

        # ãƒ©ã‚¸ã‚ªãƒœã‚¿ãƒ³ã§ç›´æ¥å…¥åŠ›ã¨ãƒ•ã‚¡ã‚¤ãƒ«ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰ã‚’é¸æŠ
        choice = st.radio("å…¥åŠ›æ–¹æ³•ã‚’é¸æŠã—ã¦ãã ã•ã„", ["ç›´æ¥å…¥åŠ›", "ãƒ•ã‚¡ã‚¤ãƒ«ã‚’ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰"])
        st.markdown('<span style="color:red">***ã€Œãƒ•ã‚¡ã‚¤ãƒ«ã‚’ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰ã€ã¯ç¤¾å†…ç’°å¢ƒã‹ã‚‰ã¯å®Ÿè¡Œã—ãªã„ã§ãã ã•ã„**</span>', unsafe_allow_html=True)

        # ç›´æ¥å…¥åŠ›ãŒé¸æŠã•ã‚ŒãŸå ´åˆ
        if choice == "ç›´æ¥å…¥åŠ›":
            user_input = st.text_area("è‡ªç”±ã«è³ªå•ã‚’å…¥åŠ›ã—ã¦ãã ã•ã„ã€‚", value=st.session_state.get("user_input_Q&A", ""), height=500)
            st.session_state["user_input_Q&A"] = user_input

        # ãƒ•ã‚¡ã‚¤ãƒ«ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰ãŒé¸æŠã•ã‚ŒãŸå ´åˆ
        elif choice == "ãƒ•ã‚¡ã‚¤ãƒ«ã‚’ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰":
            uploaded_file = st.file_uploader("ãƒ•ã‚¡ã‚¤ãƒ«ã‚’ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰", type='pdf')

            def extract_text_from_pdf(feed):
                extracted_text = ""
                with pdfplumber.open(feed) as pdf:
                    for page in pdf.pages:
                        extracted_text += page.extract_text()
                return extracted_text

            if uploaded_file is not None:
                extracted_text = extract_text_from_pdf(uploaded_file)
                user_input = st.text_area("PDFã‹ã‚‰æŠ½å‡ºã—ãŸãƒ†ã‚­ã‚¹ãƒˆ:", value=extracted_text, height=500)
                st.session_state["user_input_Q&A"] = user_input

        # ãƒ¦ãƒ¼ã‚¶ãƒ¼å…¥åŠ›ã®ç¢ºèª
        #if 'user_input' in locals() and user_input:
        #    tokens = count_tokens(user_input) - 1

        # ãƒˆãƒ¼ã‚¯ãƒ³æ•°ã‚’è¡¨ç¤º
        #    st.markdown(f'<span style="color:grey; font-size:12px;">å…¥åŠ›ã•ã‚ŒãŸãƒˆãƒ¼ã‚¯ãƒ³æ•°ï¼ˆä¸Šé™ã®ç›®å®‰ï¼š2,000ï¼‰: {tokens}</span>', unsafe_allow_html=True)
        #else:
        #    tokens = 0

        # Create a placeholder for the bot's responses
        bot_response_placeholder = st.empty()

        # Execute the communicate function when the user presses the 'Submit' button
        if st.button("å®Ÿè¡Œ", key="send_button_data"):
            if user_input.strip() == "":
                st.warning("ãƒ‡ãƒ¼ã‚¿ã‚’å…¥åŠ›ã—ã¦ãã ã•ã„ã€‚")
            else:
                st.session_state["user_input_Q&A"] = user_input
                communicate(st.session_state["user_input_Q&A"], bot_response_placeholder, model, temperature, top_p)

    # ç´¯ç©ãƒˆãƒ¼ã‚¯ãƒ³æ•°ãƒªã‚»ãƒƒãƒˆãƒœã‚¿ãƒ³ã®è¨­ç½®
    if st.sidebar.button("ãƒˆãƒ¼ã‚¯ãƒ³æ•°ãƒªã‚»ãƒƒãƒˆ"):
        st.session_state["messages"] = [
            {"role": "system", "content": "You are the best AI assistant in the world."}
        ]

        # Clear the user input
        st.session_state["user_input_Q&A"] = ""


    elif selected_option == "Translation":
        st.title("Translation")

        # ç•™æ„ç‚¹ã®è¡¨ç¤º
        st.markdown('<span style="color:red">***å€‹äººæƒ…å ±ã‚„æ©Ÿå¯†æƒ…å ±ã¯å…¥åŠ›ã—ãªã„ã§ãã ã•ã„**</span>', unsafe_allow_html=True)

        # ãƒ¦ãƒ¼ã‚¶ãƒ¼å…¥åŠ›ã‚’åˆæœŸåŒ–
        user_input = ""
        uploaded_file = ""

        # ãƒ©ã‚¸ã‚ªãƒœã‚¿ãƒ³ã§ç›´æ¥å…¥åŠ›ã¨ãƒ•ã‚¡ã‚¤ãƒ«ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰ã‚’é¸æŠ
        choice = st.radio("å…¥åŠ›æ–¹æ³•ã‚’é¸æŠã—ã¦ãã ã•ã„", ["ç›´æ¥å…¥åŠ›", "ãƒ•ã‚¡ã‚¤ãƒ«ã‚’ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰"])
        st.markdown('<span style="color:red">***ã€Œãƒ•ã‚¡ã‚¤ãƒ«ã‚’ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰ã€ã¯ç¤¾å†…ç’°å¢ƒã‹ã‚‰ã¯å®Ÿè¡Œã—ãªã„ã§ãã ã•ã„**</span>', unsafe_allow_html=True)

        # ç›´æ¥å…¥åŠ›ãŒé¸æŠã•ã‚ŒãŸå ´åˆ
        if choice == "ç›´æ¥å…¥åŠ›":
            # session_stateã®æ›´æ–°
            if "user_input_translation" in st.session_state:
                default_value = st.session_state["user_input_translation"]
            else:
                default_value = ""
            # ã‚¦ã‚£ã‚¸ã‚§ãƒƒãƒˆç”Ÿæˆ
            user_input = st.text_area("ç¿»è¨³ã—ãŸã„æ–‡ç« ã‚’å…¥åŠ›ã—ã¦ãã ã•ã„ã€‚", value=default_value, height=500, key="user_input_translation")

        # ãƒ•ã‚¡ã‚¤ãƒ«ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰ãŒé¸æŠã•ã‚ŒãŸå ´åˆ
        elif choice == "ãƒ•ã‚¡ã‚¤ãƒ«ã‚’ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰":
            uploaded_file = st.file_uploader("ãƒ•ã‚¡ã‚¤ãƒ«ã‚’ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰", type='pdf')

            def extract_text_from_pdf(feed):
                extracted_text = ""
                with pdfplumber.open(feed) as pdf:
                    for page in pdf.pages:
                        extracted_text += page.extract_text()
                return extracted_text

            if uploaded_file is not None:
                extracted_text = extract_text_from_pdf(uploaded_file)
                # session_stateã®æ›´æ–°
                st.session_state["user_input_translation"] = extracted_text
                # ã‚¦ã‚£ã‚¸ã‚§ãƒƒãƒˆç”Ÿæˆ
                user_input = st.text_area("PDFã‹ã‚‰æŠ½å‡ºã—ãŸãƒ†ã‚­ã‚¹ãƒˆ:", value=extracted_text, key="user_input_translation", height=500)


        # è¿½åŠ ï¼šè£œè¶³æƒ…å ±ã®å…¥åŠ›ãƒ•ã‚£ãƒ¼ãƒ«ãƒ‰
        additional_info = st.text_area("è£œè¶³æƒ…å ±ã‚’å…¥åŠ›ã—ã¦ãã ã•ã„ã€‚", "", key="additional_info")

        # ãƒ¦ãƒ¼ã‚¶ãƒ¼å…¥åŠ›ã®ç¢ºèª
        #if 'user_input' in locals() and user_input:
        #    tokens = count_tokens(user_input) - 2

        # ãƒˆãƒ¼ã‚¯ãƒ³æ•°ã‚’è¡¨ç¤º
        #    st.markdown(f'<span style="color:grey; font-size:12px;">å…¥åŠ›ã•ã‚ŒãŸãƒˆãƒ¼ã‚¯ãƒ³æ•°ï¼ˆä¸Šé™ã®ç›®å®‰ï¼š2,000ï¼‰: {tokens}</span>', unsafe_allow_html=True)
        #else:
        #    tokens = 0

        # Create a placeholder for the bot's responses
        bot_response_placeholder = st.empty()

        initial_prompt = (
                    "ã‚ãªãŸã¯å„ªç§€ãªç¿»è¨³å®¶ã§ã™ã€‚ã‚ãªãŸã®å½¹å‰²ã¯ã€è‹±æ–‡ã‚’æ—¥æœ¬èªã«ç¿»è¨³ã—ã€æ—¥æœ¬èªã®ã‚¦ã‚§ãƒ–ã‚µã‚¤ãƒˆä¸Šã§æ—¥æœ¬äººã®æŠ•è³‡å®¶å‘ã‘ã«ç¿»è¨³ã•ã‚ŒãŸé–“é•ã„ã®ãªã„æƒ…å ±ã‚’æä¾›ã™ã‚‹ã“ã¨ã§ã™ã€‚\n"
                    "ä»¥ä¸‹ã®æŒ‡ç¤º1ã‹ã‚‰æŒ‡ç¤º3ã«å¾“ã£ã¦ä½œæ¥­ã‚’è¡Œã£ã¦ãã ã•ã„ã€‚\n"
                    "å‡ºåŠ›ã¯ä¸‹è¨˜ã®ã€Œå½¢å¼ã€ã«å¾“ã„markdownå½¢å¼ã¨ã—ã€ã€Œ#æŒ‡ç¤ºã€ã®æ–‡è¨€ã¯å‡ºåŠ›ã—ãªã„ã§ãã ã•ã„ã€‚\n"
                    "ã€ŒPictetã€ã¨ã„ã†å˜èªã¯å¿…ãšã€Œãƒ”ã‚¯ãƒ†ã€ã¨ç¿»è¨³ã—ã¦ãã ã•ã„ã€‚\n"
                    "Wordå½¢å¼ã§ã®å‡ºåŠ›ã«ã¤ã„ã¦ã¯ã€å„ãƒ‘ãƒ©ã‚°ãƒ©ãƒ•ã‚„é …ç›®ã”ã¨ã«å¿…ãšï¼’è¡Œä»¥ä¸Šã®é–“éš”ã‚’é–‹ã‘ã¦ãã ã•ã„ã€‚\n"
                    "ï¼ƒæŒ‡ç¤º1\n"
                    f"{user_input}ã‚’ã€ä¸‹è¨˜ã®ã€Œæ³¨æ„ã—ã¦ã»ã—ã„ç‚¹ã€ã‚’å‚ç…§ã—ãªãŒã‚‰ã€å¯èƒ½ãªé™ã‚ŠåŸæ–‡ã«å¿ å®Ÿã«ã€æ¼ã‚Œã‚„é–“é•ã„ãªãã€è‡ªç„¶ãªæ—¥æœ¬èªã«ç¿»è¨³ã—ã€ã€ç¿»è¨³çµæœã€‘ã¨ã—ã¦å‡ºåŠ›ã—ã¦ãã ã•ã„\n"
                    f"ï¼ƒè£œè¶³æƒ…å ±: {additional_info}"
                    "ï¼ƒæ³¨æ„ã—ã¦ã»ã—ã„ç‚¹ï¼šæ‰€æœ‰æ ¼ã‚’ç„¡ç†ã«å…¨éƒ¨è¨³ã•ãªã„\n"
                    "ï¼ƒä¾‹â‘ \n"
                    "ã€è‹±æ–‡ã€‘At some point, our kids will be out in the world and their self-esteem will be pivotal to their success. \n"
                    "ã€æ‚ªã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘ã„ã¤ã‹ç§ãŸã¡ã®å­ä¾›ãŸã¡ãŒä¸–ç•Œã«é£›ã³ç«‹ã£ãŸæ™‚ã€å½¼ã‚‰ã®è‡ªå°Šå¿ƒã¯æˆåŠŸã®å¤§ããªè¦ã¨ãªã‚‹ã§ã—ã‚‡ã†ã€‚ \n"
                    "ã€è‰¯ã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘ã„ã¤ã‹å­ä¾›ãŸã¡ãŒä¸–ç•Œã«æ—…ç«‹ã£ãŸã¨ãã€è‡ªå°Šå¿ƒã¯æˆåŠŸã®å¤§ããªè¦ã¨ãªã‚‹ã§ã—ã‚‡ã†ã€‚\n"
                    "ï¼ƒä¾‹â‘¡\n"
                    "ã€è‹±æ–‡ã€‘The Company aims to nearly double its number of restaurants. \n"
                    "ã€æ‚ªã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘ãã®ä¼šç¤¾ã¯è‡ªç¤¾ã®ãƒ¬ã‚¹ãƒˆãƒ©ãƒ³ã®åº—èˆ—æ•°ã‚’ã»ã¼å€ã«ã™ã‚‹ã“ã¨ã‚’ç›®æŒ‡ã—ã¦ã„ã‚‹ã€‚ \n"
                    "ã€è‰¯ã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘ãã®ä¼šç¤¾ã¯ãƒ¬ã‚¹ãƒˆãƒ©ãƒ³ã®åº—èˆ—æ•°ã‚’ã»ã¼å€ã«ã™ã‚‹ã“ã¨ã‚’ç›®æŒ‡ã—ã¦ã„ã‚‹ã€‚ \n"
                    "ï¼ƒæ³¨æ„ã—ã¦ã»ã—ã„ç‚¹ï¼šè¤‡æ•°å½¢ã¯çŠ¶æ³ã«ã‚ˆã£ã¦ã¯ç„¡ç†ã«è¨³ã•ãªã„\n"
                    "ï¼ƒä¾‹â‘ \n"
                    "ã€è‹±æ–‡ã€‘The task of facilitating language learning for our children may seem complicated.\n"
                    "ã€æ‚ªã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘å­ä¾›ãŸã¡ã«å¤–å›½èªã‚’å­¦ã°ã›ã‚‹ã“ã¨ã¯é›£ã—ã„ã‚ˆã†ã«æ€ã†ã‹ã‚‚ã—ã‚Œã¾ã›ã‚“ã€‚\n"
                    "ã€è‰¯ã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘å­ä¾›ã«å¤–å›½èªã‚’å­¦ã°ã›ã‚‹ã“ã¨ã¯é›£ã—ã„ã‚ˆã†ã«æ€ã†ã‹ã‚‚ã—ã‚Œã¾ã›ã‚“ã€‚\n"
                    "ï¼ƒä¾‹â‘¡\n"
                    "ã€åŸæ–‡ã€‘For parents, preparing a list of questions before an appointment is a good start as teachers are busy.\n"
                    "ã€æ‚ªã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘æ•™å¸«ãŸã¡ã¯å¿™ã—ã„ã®ã§è¦ªã¯ã‚ã‚‰ã‹ã˜ã‚è³ªå•ã—ãŸã„ã“ã¨ã‚’æ›¸ãå‡ºã—ã¦é¢è«‡ã«è‡¨ã‚€ã¨ã„ã„ã§ã—ã‚‡ã†ã€‚\n"
                    "ã€è‰¯ã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘æ•™å¸«ã¯å¿™ã—ã„ã®ã§è¦ªã¯ã‚ã‚‰ã‹ã˜ã‚è³ªå•ã—ãŸã„ã“ã¨ã‚’æ›¸ãå‡ºã—ã¦é¢è«‡ã«è‡¨ã‚€ã¨ã„ã„ã§ã—ã‚‡ã†ã€‚ \n"
                    "ï¼ƒæ³¨æ„ã—ã¦ã»ã—ã„ç‚¹ï¼šã€Œanyã€ã¯ã€Œã‚‚ã—ï½ãªã‚‰ã€ã«åˆ†è§£ã—ãŸã»ã†ãŒã„ã„å ´åˆã‚‚ã‚ã‚‹\n"
                    "ï¼ƒä¾‹â‘ \n"
                    "ã€è‹±æ–‡ã€‘Any accident should be reported to the supervisor immediately.\n"
                    "ã€æ‚ªã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘ã©ã‚“ãªäº‹æ•…ã‚‚ä¸Šå¸ã«å ±å‘Šã•ã‚Œãªã‘ã‚Œã°ãªã‚‰ãªã„ã€‚\n"
                    "ã€è‰¯ã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘äº‹æ•…ãŒã‚ã£ãŸå ´åˆã¯å¿…ãšä¸Šå¸ã«å ±å‘Šã—ãªã‘ã‚Œã°ãªã‚‰ãªã„ã€‚\n"
                    "ï¼ƒä¾‹â‘¡\n"
                    "ã€åŸæ–‡ã€‘Any member who is in doubt should submit a copy of the medical certificate from their doctor. \n"
                    "ã€æ‚ªã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘ç–‘ã„ã®ã‚ã‚‹ã„ãšã‚Œã®ãƒ¡ãƒ³ãƒãƒ¼ã‚‚ã€åŒ»å¸«ã®è¨ºæ–­æ›¸ã‚’æå‡ºã—ãªã‘ã‚Œã°ãªã‚‰ãªã„ã€‚\n"
                    "ã€è‰¯ã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘è‡ªç„¶ãªè¨³æ–‡ï¼šãƒ¡ãƒ³ãƒãƒ¼ã¯ã€ç–‘ã„ãŒã‚ã‚‹å ´åˆã¯å¿…ãšåŒ»å¸«ã®è¨ºæ–­æ›¸ã‚’æå‡ºã—ãªã‘ã‚Œã°ãªã‚‰ãªã„ã€‚ \n"
                    "ï¼ƒæ³¨æ„ã—ã¦ã»ã—ã„ç‚¹ï¼šåè©ã‚’å‹•è©ã«ã€å‹•è©ã‚’åè©ã«å¤‰æ›ã—ãŸã»ã†ãŒè‰¯ã„å ´åˆã‚‚ã‚ã‚‹\n"
                    "ï¼ƒä¾‹â‘ ï¼šåè©å¥ã‚’å‹•è©å¥ã«å¤‰æ›ã™ã‚‹å ´åˆ\n"
                    "ã€è‹±æ–‡ã€‘Exposure to organophosphates can cause headache and diarrhea.\n"
                    "ã€æ‚ªã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘æœ‰æ©Ÿãƒªãƒ³é…¸ã¸ã®æš´éœ²ã¯é ­ç—›ã¨ä¸‹ç—¢ã‚’å¼•ãèµ·ã“ã™ã“ã¨ãŒã‚ã‚‹ã€‚\n"
                    "ã€è‰¯ã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘æœ‰æ©Ÿãƒªãƒ³é…¸ã«æ™’ã•ã‚Œã‚‹ã¨é ­ç—›ã¨ä¸‹ç—¢ãŒç”Ÿã˜ã‚‹ã“ã¨ãŒã‚ã‚‹ã€‚\n"
                    "ï¼ƒä¾‹â‘¡ï¼šå‹•è©ã‚’åè©ã«å¤‰æ›ã™ã‚‹å ´åˆã®è‹±å’Œç¿»è¨³ä¾‹\n"
                    "ã€è‹±æ–‡ã€‘The strong sales of Japanese comic books is attributable to the expansion of the international e-commerce market.\n"
                    "ã€æ‚ªã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘æ—¥æœ¬ã®ãƒãƒ³ã‚¬ã®å¥½èª¿ãªå£²ã‚Šä¸Šã’ã¯æµ·å¤–ã®ï¼¥ã‚³ãƒãƒ¼ã‚¹å¸‚å ´ã®æ‹¡å¤§ã«èµ·å› ã™ã‚‹ã€‚\n"
                    "ã€è‰¯ã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘æ—¥æœ¬ã®ãƒãƒ³ã‚¬ã®å£²ä¸ŠãŒå¥½èª¿ãªç†ç”±ã¨ã—ã¦æµ·å¤–ã®ï¼¥ã‚³ãƒãƒ¼ã‚¹å¸‚å ´ã®æ‹¡å¤§ãŒæŒ™ã’ã‚‰ã‚Œã‚‹ã€‚ \n"
                    "ï¼ƒæ³¨æ„ã—ã¦ã»ã—ã„ç‚¹ï¼šå—å‹•æ…‹ã‚’èƒ½å‹•æ…‹ã«ã€èƒ½å‹•æ…‹ã‚’å—å‹•æ…‹ã«å¤‰æ›ã—ãŸã»ã†ãŒè‰¯ã„å ´åˆã‚‚ã‚ã‚‹\n"
                    "ï¼ƒä¾‹â‘ ï¼šå—å‹•æ…‹ã‚’èƒ½å‹•æ…‹ã«å¤‰æ›ã™ã‚‹å ´åˆ\n"
                    "#â‘ â€a\n"
                    "ã€è‹±æ–‡ã€‘They wer examined by their respective family doctors.\n"
                    "ã€æ‚ªã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘å½¼ã‚‰ã¯ãã‚Œãã‚Œã‹ã‹ã‚Šã¤ã‘åŒ»ã«ã‚ˆã‚Šè¨ºå¯Ÿã•ã‚ŒãŸã€‚\n"
                    "ã€è‰¯ã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘å½¼ã‚‰ã¯ãã‚Œãã‚Œã‹ã‹ã‚Šã¤ã‘åŒ»ã®è¨ºå¯Ÿã‚’å—ã‘ãŸã€‚\n"
                    "#â‘ -b\n"
                    "ã€åŸæ–‡ã€‘Any problem has to be resolved by employees.\n"
                    "ã€æ‚ªã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘ã„ã‹ãªã‚‹å•é¡Œã‚‚å¾“æ¥­å“¡ã«ã‚ˆã£ã¦è§£æ±ºã•ã‚Œãªã‘ã‚Œã°ãªã‚‰ãªã„ã€‚\n"
                    "ã€è‰¯ã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘ã„ã‹ãªã‚‹å•é¡Œã‚‚å¾“æ¥­å“¡ãŒè§£æ±ºã—ãªã‘ã‚Œã°ãªã‚‰ãªã„ã€‚\n"
                    "ï¼ƒä¾‹â‘¡èƒ½å‹•æ…‹ã‚’å—å‹•æ…‹ã«å¤‰æ›ã™ã‚‹å ´åˆ\n"
                    "ã€è‹±æ–‡ã€‘How technology enables business model innovation.\n"
                    "ã€æ‚ªã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘ãƒ†ã‚¯ãƒãƒ­ã‚¸ãƒ¼ãŒã„ã‹ã«ãƒ“ã‚¸ãƒã‚¹ãƒ¢ãƒ‡ãƒ«ã®ã‚¤ãƒãƒ™ãƒ¼ã‚·ãƒ§ãƒ³ã‚’å¯èƒ½ã«ã—ã¦ã„ã‚‹ã‹ã€‚\n"
                    "ã€è‰¯ã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘ãƒ†ã‚¯ãƒãƒ­ã‚¸ãƒ¼ã«ã‚ˆã‚Šã„ã‹ã«ãƒ“ã‚¸ãƒã‚¹ãƒ¢ãƒ‡ãƒ«ã®ã‚¤ãƒãƒ™ãƒ¼ã‚·ãƒ§ãƒ³ãŒã‚‚ãŸã‚‰ã•ã‚Œã‚‹ã‹ã€‚ \n"
                    "ï¼ƒæ³¨æ„ã—ã¦ã»ã—ã„ç‚¹ï¼šä½¿å½¹å‹•è©ã¯ã‹ã¿ç •ã„ã¦è¨³ã—ãŸæ–¹ãŒã„ã„å ´åˆãŒå¤šã„\n"
                    "ï¼ƒä¾‹â‘ \n"
                    "ã€è‹±æ–‡ã€‘This combination of experience and innovation has made the company so successful. \n"
                    "ã€æ‚ªã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘ã“ã®çµŒé¨“ã¨ã‚¤ãƒãƒ™ãƒ¼ã‚·ãƒ§ãƒ³ã®çµ„ã¿åˆã‚ã›ãŒãã®ä¼æ¥­ã‚’æˆåŠŸã•ã›ãŸã€‚ \n"
                    "ã€è‰¯ã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘ã“ã®çµŒé¨“ã¨ã‚¤ãƒãƒ™ãƒ¼ã‚·ãƒ§ãƒ³ã“ããŒãã®ä¼æ¥­ã‚’æˆåŠŸã«å°ã„ãŸè¦å› ã ã€‚\n"
                    "ï¼ƒä¾‹â‘¡\n"
                    "ã€åŸæ–‡ã€‘Professor Smith has made me want to become a teacher.\n"
                    "ã€æ‚ªã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘ã‚¹ãƒŸã‚¹æ•™æˆã¯ç§ã‚’å…ˆç”Ÿã«ãªã‚ŠãŸãã•ã›ãŸã€‚\n"
                    "ã€è‰¯ã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘ã‚¹ãƒŸã‚¹æ•™æˆã«å‡ºä¼šã£ã¦ç§ã¯å…ˆç”Ÿã«ãªã‚ŠãŸã„ã¨æ€ã£ãŸã€‚\n"
                    "ï¼ƒæ³¨æ„ã—ã¦ã»ã—ã„ç‚¹ï¼šã€Œï½ãŸã‚ã®ã€ã®ã€Œtoã€ã‚„ã€Œforã€ã‚’è¨³ã—ä¸‹ã’ã‚‹\n"
                    "ï¼ƒä¾‹â‘ \n"
                    "ã€è‹±æ–‡ã€‘Lisa had turned her head to observe the birds climbing into the blue sky. \n"
                    "ã€æ‚ªã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘ãƒªã‚µã¯é³¥ãŸã¡ãŒé’ã„ç©ºã¸ã¨é£›ã³ç«‹ã£ã¦ã„ãã®ã‚’è¦‹ã‚‹ãŸã‚ã«æŒ¯ã‚Šè¿”ã£ãŸã€‚\n"
                    "ã€è‰¯ã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘ãƒªã‚µãŒæŒ¯ã‚Šè¿”ã‚‹ã¨é³¥ãŸã¡ãŒé’ã„ç©ºã¸ã¨é£›ã³ç«‹ã£ã¦ã„ãã®ãŒè¦‹ãˆãŸã€‚\n"
                    "ï¼ƒä¾‹â‘¡\n"
                    "ã€è‹±æ–‡ã€‘The application shall be submitted to the president for review. \n"
                    "ã€æ‚ªã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘ç”³è¾¼æ›¸ã¯ç¢ºèªã®ãŸã‚ã«ç¤¾é•·ã«æå‡ºã•ã‚Œãªã‘ã‚Œã°ãªã‚‰ãªã„ã€‚\n"
                    "ã€è‰¯ã„æ—¥æœ¬èªè¨³ã®ä¾‹ã€‘ç”³è¾¼æ›¸ã‚’æå‡ºã—ç¤¾é•·ã®ç¢ºèªã‚’å—ã‘ãªã‘ã‚Œã°ãªã‚‰ãªã„ã€‚\n"
                    "###\n"
                    "ï¼ƒæŒ‡ç¤º2\n"
                    "#æŒ‡ç¤º1ã§ç¿»è¨³ã«ã‚ˆã‚Šä½œæˆã•ã‚ŒãŸæ–‡ç« ã‚’ã€åŠåˆ†ã®åˆ†é‡ã«ãªã‚‹ã‚ˆã†è¦ç´„ã—ã€ã€è¦ç´„ã€‘ã¨ã—ã¦å‡ºåŠ›ã—ã¦ãã ã•ã„ã€‚\n"
                    "###\n"
                    "ï¼ƒæŒ‡ç¤º3\n"
                    "#æŒ‡ç¤º1ã§ç¿»è¨³ã«ã‚ˆã‚Šä½œæˆã•ã‚ŒãŸæ–‡ç« ã‹ã‚‰ã€å›ºæœ‰åè©ã‚’æŠœãå‡ºã—ã¦ãƒªã‚¹ãƒˆåŒ–ã—ã€ç°¡å˜ãªèª¬æ˜ã‚’ã¤ã‘ã¦ã€ã€å›ºæœ‰åè©ã¨ãã®èª¬æ˜ã€‘ã¨ã—ã¦å‡ºåŠ›ã—ã¦ãã ã•ã„ã€‚\n"
                    "###\n"
        )
        # ç¿»è¨³ã®å®Ÿè¡Œã‚³ãƒãƒ³ãƒ‰
        if st.button("å®Ÿè¡Œ", key="send_button_translation"):
            if user_input.strip() == "":
                st.warning("ãƒ‡ãƒ¼ã‚¿ã‚’å…¥åŠ›ã—ã¦ãã ã•ã„ã€‚")
            else:
                # æ–°ã—ã„ã‚»ãƒƒã‚·ãƒ§ãƒ³ã”ã¨ã«ãƒ¡ãƒƒã‚»ãƒ¼ã‚¸å±¥æ­´ã‚’ãƒªã‚»ãƒƒãƒˆ
                st.session_state["messages"] = []
                # ã‚¤ãƒ‹ã‚·ãƒ£ãƒ«ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã®å…¥åŠ›ã¨ãƒãƒ£ãƒƒãƒˆã®ç”Ÿæˆ
                st.session_state["user_input"] = initial_prompt
                generated_text = communicate(initial_prompt, bot_response_placeholder, model, temperature, top_p)

                # ç”Ÿæˆã•ã‚ŒãŸãƒ†ã‚­ã‚¹ãƒˆã‚’UIã«è¡¨ç¤ºã—ã¾ã™ã€‚
                #bot_response_placeholder = st.write(generated_text)



                # Wordæ–‡æ›¸ã‚’ç”Ÿæˆ
                doc_path = create_word_doc(generated_text)

                def get_binary_file_downloader_html(bin_file, file_label='File'):
                    with open(bin_file, 'rb') as f:
                        data = f.read()
                    bin_str = base64.b64encode(data).decode()
                    href = f'<a href="data:application/octet-stream;base64,{bin_str}" download="{bin_file}">{file_label}</a>'
                    return href
                    # ãƒ•ã‚¡ã‚¤ãƒ«åã‚’åŸæ–‡ã‚¿ã‚¤ãƒˆãƒ«ã‹ã‚‰å–å¾—ã™ã‚‹ãƒ‘ã‚¿ãƒ¼ãƒ³ï¼ˆä¿®æ­£ä¸­ï¼‰
                    #href = f'<a href="data:application/octet-stream;base64,{bin_str}" download="{os.path.basename(bin_file)}">{file_label}</a>'

                # Wordæ–‡æ›¸ã‚’ãƒ€ã‚¦ãƒ³ãƒ­ãƒ¼ãƒ‰ãƒªãƒ³ã‚¯ã¨ã—ã¦æä¾›
                st.markdown(get_binary_file_downloader_html(doc_path, "çµæœã‚’Wordå½¢å¼ã§ãƒ€ã‚¦ãƒ³ãƒ­ãƒ¼ãƒ‰"), unsafe_allow_html=True)
                # ãƒ•ã‚¡ã‚¤ãƒ«åã‚’åŸæ–‡ã‚¿ã‚¤ãƒˆãƒ«ã‹ã‚‰å–å¾—ã™ã‚‹ãƒ‘ã‚¿ãƒ¼ãƒ³ï¼ˆä¿®æ­£ä¸­ï¼‰
                #st.markdown(get_binary_file_downloader_html(doc_path, os.path.basename(doc_path)), unsafe_allow_html=True)

                # ãƒ¡ãƒƒã‚»ãƒ¼ã‚¸å±¥æ­´ã®åˆæœŸåŒ–
                if "messages" not in st.session_state:
                    st.session_state["messages"] = []

        # APIã«é€ä¿¡ã™ã‚‹ãƒ‡ãƒ¼ã‚¿ã‚’è¡¨ç¤ºã™ã‚‹å‰ã«ã€`messages` å¤‰æ•°ã®çŠ¶æ…‹ã‚’ç¢ºèª
        #if "messages" in st.session_state:
        #    messages = st.session_state["messages"]
        #    st.write("é€ä¿¡ã™ã‚‹ãƒªã‚¯ã‚¨ã‚¹ãƒˆ:", {"model": model, "messages": messages, "temperature": temperature, "top_p": top_p})
        #else:
        #    st.write("ãƒ¡ãƒƒã‚»ãƒ¼ã‚¸ãŒæœªå®šç¾©ã§ã™ã€‚")


        # ã€Œã‚·ã‚¹ãƒ†ãƒ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã‚’è¡¨ç¤ºã€ãƒœã‚¿ãƒ³ã®èª¬æ˜
        st.markdown('<span style="color:grey; font-size:12px;">***ä¸‹ã®ã€Œã‚·ã‚¹ãƒ†ãƒ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã‚’è¡¨ç¤ºã€ãƒœã‚¿ãƒ³ã‚’æŠ¼ã™ã¨ã€ã“ã®æ©Ÿèƒ½ã«ã‚ã‚‰ã‹ã˜ã‚çµ„ã¿è¾¼ã¾ã‚Œã¦ã„ã‚‹ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆï¼ˆå‘½ä»¤æ–‡ï¼‰ã‚’è¡¨ç¤ºã§ãã¾ã™ã€‚**</span>', unsafe_allow_html=True)

        # ã€Œã‚·ã‚¹ãƒ†ãƒ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã‚’è¡¨ç¤ºã€ãƒœã‚¿ãƒ³ã®è¨­ç½®
        if st.button("ã‚·ã‚¹ãƒ†ãƒ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã‚’è¡¨ç¤º"):
            st.write(initial_prompt)


    elif selected_option == "Proofreading":
        st.title("Proofreading")

        # ç•™æ„ç‚¹ã®è¡¨ç¤º
        st.markdown('<span style="color:red">***å€‹äººæƒ…å ±ã‚„æ©Ÿå¯†æƒ…å ±ã¯å…¥åŠ›ã—ãªã„ã§ãã ã•ã„**</span>', unsafe_allow_html=True)

        # ãƒ¦ãƒ¼ã‚¶ãƒ¼å…¥åŠ›ã‚’åˆæœŸåŒ–
        user_input = ""
        uploaded_file = ""

        # ãƒ©ã‚¸ã‚ªãƒœã‚¿ãƒ³ã§ç›´æ¥å…¥åŠ›ã¨ãƒ•ã‚¡ã‚¤ãƒ«ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰ã‚’é¸æŠ
        choice = st.radio("å…¥åŠ›æ–¹æ³•ã‚’é¸æŠã—ã¦ãã ã•ã„", ["ç›´æ¥å…¥åŠ›", "ãƒ•ã‚¡ã‚¤ãƒ«ã‚’ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰"])
        st.markdown('<span style="color:red">***ã€Œãƒ•ã‚¡ã‚¤ãƒ«ã‚’ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰ã€ã¯ç¤¾å†…ç’°å¢ƒã‹ã‚‰ã¯å®Ÿè¡Œã—ãªã„ã§ãã ã•ã„**</span>', unsafe_allow_html=True)

        # ç›´æ¥å…¥åŠ›ãŒé¸æŠã•ã‚ŒãŸå ´åˆ
        if choice == "ç›´æ¥å…¥åŠ›":
            # session_stateã®æ›´æ–°
            if "user_input_proof" in st.session_state:
                default_value = st.session_state["user_input_proof"]
            else:
                default_value = ""
            # ã‚¦ã‚£ã‚¸ã‚§ãƒƒãƒˆç”Ÿæˆ
            user_input = st.text_area("æ ¡é–²/æ ¡æ­£ã—ãŸã„æ–‡ç« ã‚’å…¥åŠ›ã—ã¦ãã ã•ã„ã€‚", value=default_value, height=500, key="user_input_proof")

        # ãƒ•ã‚¡ã‚¤ãƒ«ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰ãŒé¸æŠã•ã‚ŒãŸå ´åˆ
        elif choice == "ãƒ•ã‚¡ã‚¤ãƒ«ã‚’ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰":
            uploaded_file = st.file_uploader("ãƒ•ã‚¡ã‚¤ãƒ«ã‚’ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰", type='pdf')

            def extract_text_from_pdf(feed):
                extracted_text = ""
                with pdfplumber.open(feed) as pdf:
                    for page in pdf.pages:
                        extracted_text += page.extract_text()
                return extracted_text

            if uploaded_file is not None:
                extracted_text = extract_text_from_pdf(uploaded_file)
                # session_stateã®æ›´æ–°
                st.session_state["user_input_proof"] = extracted_text
                # ã‚¦ã‚£ã‚¸ã‚§ãƒƒãƒˆç”Ÿæˆ
                user_input = st.text_area("PDFã‹ã‚‰æŠ½å‡ºã—ãŸãƒ†ã‚­ã‚¹ãƒˆ:", value=extracted_text, key="user_input_proof", height=500)

        # è¿½åŠ ï¼šè£œè¶³æƒ…å ±ã®å…¥åŠ›ãƒ•ã‚£ãƒ¼ãƒ«ãƒ‰
        additional_info = st.text_area("è£œè¶³æƒ…å ±ã‚’å…¥åŠ›ã—ã¦ãã ã•ã„ã€‚", "", key="additional_info")

        # ãƒ¦ãƒ¼ã‚¶ãƒ¼ã®å…¥åŠ›ã‚’è¡¨ç¤º
        #st.write("ãƒ¦ãƒ¼ã‚¶ãƒ¼å…¥åŠ›:", user_input)

        # ãƒ¦ãƒ¼ã‚¶ãƒ¼å…¥åŠ›ã®ç¢ºèª
        #if 'user_input' in locals() and user_input:
        #    tokens = count_tokens(user_input) - 2

        # ãƒˆãƒ¼ã‚¯ãƒ³æ•°ã‚’è¡¨ç¤º
        #    st.markdown(f'<span style="color:grey; font-size:12px;">å…¥åŠ›ã•ã‚ŒãŸãƒˆãƒ¼ã‚¯ãƒ³æ•°ï¼ˆä¸Šé™ã®ç›®å®‰ï¼š2,000ï¼‰: {tokens}</span>', unsafe_allow_html=True)
        #else:
        #    tokens = 0

        # Create a placeholder for the bot's responses
        bot_response_placeholder = st.empty()

        initial_prompt = (
            "## ã‚ãªãŸã¯æ ¡é–²ãƒ»æ ¡æ­£ã®å„ªç§€ãªã‚¹ãƒšã‚·ãƒ£ãƒªã‚¹ãƒˆã§ã™ã€‚  \n"
            "ã‚ãªãŸã®å½¹å‰²ã¯ã€æ—¥æœ¬ã®æŠ•è³‡å®¶å‘ã‘ã«å…¬è¡¨ã•ã‚Œã‚‹æƒ…å ±ã‚’æ ¡é–²ãƒ»æ ¡æ­£ã—ã€é–“é•ã„ãªãé«˜å“è³ªãªæ–‡ç« ã‚’ä½œæˆã™ã‚‹ã“ã¨ã§ã™ã€‚  \n"
            "ã“ã‚Œã‹ã‚‰å…¥åŠ›ã™ã‚‹æ–‡ç« ã«å¯¾ã—ã¦ã€ä¸‹è¨˜ã®å‡¦ç†ã¨å‡ºåŠ›ã‚’è¨˜è¼‰ã•ã‚ŒãŸé †ç•ªé€šã‚Šè¡Œã£ã¦ãã ã•ã„ã€‚  \n"
            "ãªãŠã€å‡¦ç†1ã¨å‡¦ç†2ã«ã¤ã„ã¦ã¯ã€å¾Œã®å‡ºåŠ›1ãƒ»2ã®ãŸã‚ã®å‡¦ç†ã‚’è¡Œã†ã®ã¿ã¨ã—ã€çµæœã‚’æ–‡ç« ã¨ã—ã¦è¡¨ç¤ºã—ãªã„ã§ãã ã•ã„ã€‚  \n"
            "è¡¨ç¤ºã¯å‡ºåŠ›1ã®ã€Œä¿®æ­£å¾Œå…¨æ–‡ã€ã¨å‡ºåŠ›2ã€Œä¿®æ­£ç®‡æ‰€ãƒªã‚¹ãƒˆã€ã€åŠã³å‡ºåŠ›3ã€Œè¡¨è¨˜æºã‚Œã¨è€ƒãˆã‚‰ã‚Œã‚‹ã‚‚ã®ã€ã«ãŠã„ã¦ã®ã¿è¡Œã£ã¦ãã ã•ã„ã€‚  \n"
            "ã¾ãŸã€è©²å½“ã™ã‚‹ã‚‚ã®ãŒãªã„é …ç›®ã«ã¤ã„ã¦ã‚‚ã€è¡¨ç¤ºã¯ã—ãªã„ã§ãã ã•ã„ \n"
            "å…¥åŠ›ã•ã‚ŒãŸæ–‡ç« ã®æ–‡ã®é€”ä¸­ã«æ”¹è¡ŒãŒå…¥ã£ã¦ã„ãŸå ´åˆã¯ã€é©å®œæ”¹è¡Œã‚’å‰Šé™¤ã—ã¦å‡¦ç†ã‚’è¡Œã£ã¦ãã ã•ã„ã€‚\n"
            "###  \n"
            "å‡¦ç†1:èª¤ã‚Šã®æ¤œçŸ¥ï¼ˆè¡¨ç¤ºã—ãªã„ï¼‰ \n"
            "ä»¥ä¸‹ã®åŸºæº–ã«è©²å½“ã™ã‚‹ã‚‚ã®ã‚’ã™ã¹ã¦æ¤œçŸ¥ã—ã¦ãã ã•ã„ã€‚ \n"
            "æ¤œçŸ¥ã•ã‚ŒãŸå€‹æ‰€ã¯ã€ã©ã‚“ãªã«å¤§é‡ã«ãªã£ã¦ã‚‚æ§‹ã„ã¾ã›ã‚“ã€‚ \n"
            "æ­£ã—ã„ã‹ã©ã†ã‹ã®åˆ¤æ–­ãŒå¾®å¦™ãªã€å¾®ç´°ãªèª¤ã‚Šã¨æ€ã‚ã‚Œã‚‹ã‚‚ã®ã‚‚å«ã‚ã€å°‘ã—ã§ã‚‚é•å’Œæ„ŸãŒã‚ã‚‹ã‚‚ã®ã¯ã€ç©æ¥µçš„ã«ã§ãã‚‹ã ã‘å¤šãæ¤œçŸ¥ã—ã€ä¸€ã¤ã®æ¼ã‚Œã‚‚ãªãã€ã™ã¹ã¦ãƒªã‚¹ãƒˆåŒ–ã—ã¦ãã ã•ã„ã€‚ \n"
            "### \n"
            "åŸºæº–:  \n"
            "#1:èª¤å­—è„±å­—ã€ã‚¿ã‚¤ãƒ—ãƒŸã‚¹ã¨æ€ã‚ã‚Œã‚‹ã‚‚ã®  \n"
            "ã€ä¾‹ã€‘ \n"
            "ãƒ»Ã—ã€Œã“ã‚“ã«ã¡ã‚ã€â†’ã€‡ã€Œã“ã‚“ã«ã¡ã¯ã€ \n"
            "ãƒ»Ã—ã€Œã‚ˆã‚ã—ããŠé¡˜ã„ã—mã™ã€â†’ã€‡ã€Œã‚ˆã‚ã—ããŠé¡˜ã„ã—ã¾ã™ã€ \n"
            "ãƒ»Ã—ã€Œã“ã˜ã‚“ã¾ã‚Šã€â†’ã€‡ã€Œã“ã¢ã‚“ã¾ã‚Šã€ \n"
            "ãƒ»Ã—ã€Œã¡ã˜ã“ã¾ã‚‹ã€â†’ã€‡ã€Œã¡ã¢ã“ã¾ã‚‹ã€ \n"
            "ãƒ»Ã—ã€Œã„ã¡ã¢ã‚‹ã—ã„ã€â†’ã€‡ã€Œã„ã¡ã˜ã‚‹ã—ã„ã€ \n"
            "ãƒ»Ã—ã€Œã‹ãŸãšã‘ã‚‹ã€â†’ã€‡ã€Œã‹ãŸã¥ã‘ã‚‹ã€ \n"
            "ãƒ»Ã—ã€Œé‡‘ãšã‚‹ã€â†’ã€‡ã€Œé‡‘ã¥ã‚‹ã€ \n"
            "ãƒ»Ã—ã€Œã¤ããšãã€â†’ã€‡ã€Œã¤ãã¥ãã€ \n"
            "ãƒ»Ã—ã€Œã†ãªã¥ãã€â†’ã€‡ã€Œã†ãªãšãã€ \n"
            "ãƒ»Ã—ã€Œå£ã¥ã•ã‚€ã€â†’ã€‡ã€Œå£ãšã•ã‚€ã€ \n"
            "ãƒ»ã€Œæ™‚å­£ã€ã€Œæ™‚æœŸã€ã€Œæ™‚æ©Ÿã€ãªã©ã€æ„å‘³ãŒä¼¼é€šã£ã¦ã„ã¦ã€è¦‹ãŸç›®ã«ã‚‚åˆ¤æ–­ã—ã¥ã‚‰ãèª¤å¤‰æ›ã—ã‚„ã™ã„å˜èª \n"
            "#2:æ•°å­—ã®è¡¨è¨˜ãŒå…¨è§’ã«ãªã£ã¦ã„ã‚‹ã‚‚ã®ï¼ˆæ•°å­—ã®è¡¨è¨˜ã¯å…¨ã¦åŠè§’ã§çµ±ä¸€ã™ã‚‹ã€‚ï¼‰ \n"
            "#3:æ…£ç”¨å¥ã‚„ã“ã¨ã‚ã–ã®è¡¨ç¾ã«èª¤ã‚ŠãŒã‚ã‚‹ã¨è€ƒãˆã‚‰ã‚Œã‚‹ã‚‚ã® \n"
            "ã€ä¾‹ã€‘ \n"
            "ãƒ»Ã—ã€Œçš„ã‚’å¾—ã‚‹ã€â†’ã€‡ã€Œçš„ã‚’å°„ã‚‹ã€ \n"
            "ãƒ»Ã—ã€Œã¨ã‚“ã§ã‚‚ã‚ã‚Šã¾ã›ã‚“ã€â†’ã€‡ã€Œã¨ã‚“ã§ã‚‚ãªã„ã“ã¨ã§ã™ã€ï¼ˆã€Œã¨ã‚“ã§ã‚‚ãªã„ã€ã¯å½¢å®¹è©ã®ãŸã‚æ´»ç”¨ã—ãªã„ã€‚ï¼‰ \n"
            "#4:çµŒæ¸ˆé–¢ä¿‚ç”¨èªã«ã¤ã„ã¦ä¸‹è¨˜ã®ãƒ«ãƒ¼ãƒ«ã«æ²¿ã£ã¦ãŠã‚‰ãšèª¤ã‚ŠãŒã‚ã‚‹ã¨è€ƒãˆã‚‰ã‚Œã‚‹ã‚‚ã® \n"
            "ã€ãƒ«ãƒ¼ãƒ«ã€‘ \n"
            "ï¼ˆ1ï¼‰ä¸‹è¨˜ã®ã‚ˆã†ãªçµŒæ¸ˆé–¢ä¿‚ç”¨èªã¯ã€å˜ä½“ã§ã¯é€ã‚Šä»®åãŒå¿…è¦ \n"
            "ã€ä¾‹ã€‘ \n"
            "ãƒ»é å…¥ã‚Œ \n"
            "ãƒ»å£²ä¸Šã’ \n"
            "ãƒ»å¸å£²ã‚Š \n"
            "ãƒ»è²·æ›ã‘ \n"
            "ãƒ»åˆ©å›ã‚Š \n"
            "ï¼ˆ2ï¼‰ï¼ˆ1ï¼‰ã«ã‚ã‚‹ã‚ˆã†ãªèªãŒè¤‡åˆèªã¨ãªã‚‹ã¨ã€é€ã‚Šä»®åã¯ä¸è¦ã¨ãªã‚‹ \n"
            "ã€ä¾‹ã€‘ \n"
            "ãƒ»Ã—å—ä»˜ã‘ä»¶æ•°â†’ã€‡å—ä»˜ä»¶æ•° \n"
            "ãƒ»Ã—å£²ä¸Šã’é«˜â†’ã€‡å£²ä¸Šé«˜ \n"
            "ãƒ»Ã—å¸å£²ã‚Šå•å±‹â†’ã€‡å¸å£²å•å±‹ \n"
            "ãƒ»Ã—æ”¯æ‰•ã„ç·é¡â†’ã€‡æ”¯æ‰•ç·é¡ \n"
            "ãƒ»Ã—æŒ¯è¾¼ã¿æ‰‹æ•°æ–™â†’ã€‡æŒ¯è¾¼ã¿æ‰‹æ•°æ–™ \n"
            "ãƒ»Ã—ç”³ã—è¾¼ã¿æœŸé–“â†’ã€‡ç”³è¾¼æœŸé–“ \n"
            "ãƒ»Ã—å‰²ã‚Šå½“ã¦æ•°é‡â†’ã€‡å‰²å½“æ•°é‡ \n"
            "ãƒ»Ã—å‰²ã‚Šå¢—ã—æ–™é‡‘â†’ã€‡å‰²å¢—æ–™é‡‘ \n"
            "#5:æ–‡è„ˆã«åˆã‚ãªã„å˜èªãŒä½¿ã‚ã‚Œã¦ã„ã‚‹ã‚‚ã® \n"
            "ã€ä¾‹ã€‘ \n"
            "ãƒ»Ã—ã€ŒPERãŒæ‹¡å¤§ã—ã¦ã„ã‚‹ã€â†’ã€‡ã€ŒPERãŒä¸Šæ˜‡ã—ã¦ã„ã‚‹ã€ï¼ˆPERã¯ã€Œç‡ã€ã§ã‚ã‚‹ã®ã§æ‹¡å¤§ãƒ»ç¸®å°ã™ã‚‹ã‚‚ã®ã§ã¯ãªãä¸Šæ˜‡ãƒ»ä¸‹è½ã™ã‚‹ã‚‚ã®ï¼‰ \n"
            "ãƒ»Ã—ã€Œæ–°è¦å—æ³¨ãŒå¤§ããä¸Šæ˜‡ã™ã‚‹ã€â†’ã€‡ã€Œæ–°è¦å—æ³¨ãŒå¤§å¹…ã«å¢—åŠ ã™ã‚‹ã€ï¼ˆæ–°è¦å—æ³¨ã¯ã€Œä»¶æ•°ã€ã§ã‚ã‚‹ã®ã§ã€ä¸Šæ˜‡ãƒ»ä¸‹è½ã™ã‚‹ã‚‚ã®ã§ã¯ãªãå¢—åŠ ãƒ»æ¸›å°‘ã™ã‚‹ã‚‚ã®ï¼‰ \n"
            "ãƒ»Ã—ã€Œãƒˆãƒ¬ãƒ³ãƒ‰æŒ‡æ¨™ãŒå¼·ãæ”¹å–„ã—ã¦ã„ã‚‹ã€â†’ã€‡ã€Œãƒˆãƒ¬ãƒ³ãƒ‰æŒ‡æ¨™ãŒå¤§å¹…ã«æ”¹å–„ã—ã¦ã„ã‚‹ã€ \n"
            "ãƒ»Ã—ã€Œä¸­ç«‹é‡‘åˆ©ã¯æ¦‚ã­æ¸›å°‘å‚¾å‘ã§ã‚ã£ãŸã€â†’â—‹ã€Œä½ä¸‹å‚¾å‘ã§ã‚ã£ãŸã€  \n"
            "#6:ä¸»èªã¨è¿°èªã®çµ„ã¿åˆã‚ã›ãŒé–“é•ã£ã¦ã„ã‚‹ã‚‚ã®ã€ã¾ãŸã¯ä¸æ˜ç­ãªã‚‚ã® \n"
            "#7:æ–‡æœ«ã®è¡¨ç¾ãŒã€Œã§ã™ã€ã¾ã™ã€å£èª¿ã«ãªã£ã¦ã„ã‚‹ã‚‚ã® \n"
            "#8:å¥èª­ç‚¹ã®æ‰“ã¡æ–¹ã«ä¸è‡ªç„¶ãªç‚¹ãŒã‚ã‚‹ã‚‚ã® \n"
            "ã€ä¾‹ã€‘ \n"
            "ãƒ»å¥ç‚¹ã‚„è¨˜å·ä»¥å¤–ã§æ”¹è¡Œã—ã¦ã„ã‚‹ \n"
            "ãƒ»ä¸€æ–‡ã«èª­ç‚¹ãŒ4ã¤ä»¥ä¸Šã‚ã‚‹ \n"
            "ãƒ»50æ–‡å­—ä»¥ä¸Šã®æ–‡ã«èª­ç‚¹ãŒãªã„ \n"
            "ãƒ»ä¸€æ–‡ãŒ100æ–‡å­—ä»¥ä¸Šã‚ã‚‹ \n"
            "#9:æ‹¬å¼§ã‚„ã‚«ã‚®æ‹¬å¼§ã€ã‚¯ã‚©ãƒ¼ãƒ†ãƒ¼ã‚·ãƒ§ãƒ³ãƒãƒ¼ã‚¯ç­‰ã®å§‹ç‚¹ãƒ»çµ‚ç‚¹ãŒæ¬ ã‘ã¦ã„ã‚‹ã‚‚ã® \n"
            "#10:è©±ã—è¨€è‘‰ç­‰ã®æ­£å¼ãªæ–‡ç« ã«ãµã•ã‚ã—ããªã„è¡¨ç¾ã«ãªã£ã¦ã„ã‚‹ã‚‚ã® \n"
            "ã€ä¾‹ã€‘ \n"
            "ãƒ»ã€Œã¡ã‚ƒã‚“ã¨ã€ã‚„ã€Œã¡ã‚‡ã£ã¨ã€ \n"
            "ãƒ»ã€Œé£Ÿã¹ã‚Œãªã„ã€ï¼ˆã€Œé£Ÿã¹ã‚‰ã‚Œãªã„ã€ã®ã€Œã‚‰ã€ãŒæŠœã‘ãŸã€Œã‚‰æŠœãè¨€è‘‰ã€ï¼‰ \n"
            "ãƒ»ã€Œæ€’ã£ã¦ã‚‹ã€ï¼ˆã€Œæ€’ã£ã¦ã„ã‚‹ã€ã®ã€Œã„ã€ãŒæŠœã‘ãŸã€Œã„æŠœãè¨€è‘‰ã€ï¼‰ \n"
            "ãƒ»ã€Œèª­ã¾ã•ã›ã‚‹ã€ã‚„ã€Œèã‹ã•ã›ã‚‹ã€ï¼ˆã€Œèª­ã¾ã›ã‚‹ã€ã€Œèã‹ã›ã‚‹ã€ã«æœ¬æ¥ä¸è¦ãªã€Œã•ã€ãŒå…¥ã£ãŸã€Œã•å…¥ã‚Œè¡¨ç¾ã€ï¼‰ \n"
            "#11:é‡è¤‡è¡¨ç¾ã«ãªã£ã¦ã„ã‚‹ã‚‚ã® \n"
            "ã€ä¾‹ã€‘ \n"
            "ãƒ»ã€Œä¸€ç•ªæœ€åˆã€ \n"
            "ãƒ»ã€Œå¾Œã§å¾Œæ‚”ã™ã‚‹ã€ \n"
            "ãƒ»ã€Œå¿…ãšå¿…è¦ã€ \n"
            "ãƒ»ã€Œè¢«å®³ã‚’è¢«ã‚‹ã€ \n"
            "#12:å·®åˆ¥èªãƒ»ä¸å¿«èªï½¤ãã®ä»–ãƒãƒªãƒ†ã‚£ã‚«ãƒ«ãƒ»ã‚³ãƒ¬ã‚¯ãƒˆãƒã‚¹ã«åã™ã‚‹ã‚‚ã® \n"
            "ãƒ»ã€ŒäºŒäººä¸‰è„šã€ã‚„ã€Œæ°—é•ã„ã€ãªã©ã€ã§ãã‚Œã°ä½¿ç”¨ã™ã¹ãã§ãªã„å·®åˆ¥èªã¨ãªã‚Šã†ã‚‹è¨€è‘‰ã‚„ã€å—ã‘å–ã‚‹äººã«ã¨ã£ã¦ã¯ä¸å¿«ãªè¨€è‘‰ \n"
            "ãƒ»ã€Œãƒ“ã‚¸ãƒã‚¹ãƒãƒ³ã€ã‚„ã€Œçœ‹è­·å©¦ã€ãªã©ã€åè¦‹ã‚’å«ã‚€ã‹ã€å…¬å¹³ã§ãªã„è¨€è‘‰ã¨æ‰ãˆã‚‰ã‚Œã‚‹å¯èƒ½æ€§ã®ã‚ã‚‹è¨€è‘‰ \n"
            "#13:åŠ©è©ã®ä½¿ç”¨ãŒèª¤ã£ã¦ã„ãŸã‚Šä¸é©åˆ‡ã§ã‚ã‚‹ã¨è€ƒãˆã‚‰ã‚Œã‚‹ã‚‚ã® \n"
            "ã€ä¾‹ã€‘ \n"
            "ãƒ»Ã—ã€Œé‹ç”¨ã‚’æŸ”è»ŸåŒ–ã‚’æ±ºå®šã—ãŸå¾Œã‚‚ã€â†’ã€‡ã€Œé‹ç”¨ã®æŸ”è»ŸåŒ–ã‚’æ±ºå®šã—ãŸå¾Œã‚‚ã€ï¼ˆæ ¼åŠ©è©ã€Œã‚’ã€ã®ä½¿ç”¨ãŒä¸é©åˆ‡ï¼‰ \n"
            "ãƒ»Ã—ã€Œè‡ªç„¶åˆ©ç‡ãŒä¸Šæ˜‡ã«ã¤ã„ã¦ã¯ã€â†’ã€‡ã€Œè‡ªç„¶åˆ©å­ç‡ã®ä¸Šæ˜‡ã«ã¤ã„ã¦ã¯ã€ï¼ˆæ ¼åŠ©è©ã€ŒãŒã€ã®ä½¿ç”¨ãŒä¸é©åˆ‡ï¼‰ \n"
            "ãƒ»Ã—ã€ŒFOMCãŒé•·æœŸã®æ”¿ç­–é‡‘åˆ©ãŒã€â†’ã€‡ã€ŒFOMCã«ã‚ˆã‚‹é•·æœŸã®æ”¿ç­–é‡‘åˆ©ãŒã€ï¼ˆæ ¼åŠ©è©ã€ŒãŒã€ã®ä½¿ç”¨ãŒä¸é©åˆ‡ï¼‰ \n"
            "ãƒ»Ã—ã€Œã“ã®ãƒ–ãƒ­ã‚°å…¬è¡¨ã•ã‚Œã‚‹å‰ã€â†’ã€‡ã€Œã“ã®ãƒ–ãƒ­ã‚°ãŒå…¬è¡¨ã•ã‚Œã‚‹å‰ã€ï¼ˆæ ¼åŠ©è©ã€ŒãŒã€ã®æŠœã‘ï¼‰ \n"
            "#14:ãã®ä»–ä¸é©åˆ‡ãªè¡¨ç¾ã®ä¾‹ \n"
            "ã€ä¾‹ã€‘ \n"
            "ãƒ»Ã—ã€Œè¬›æ¼”ã®æ®‹ã‚Šã«ã¯ã€â†’ã€‡ã€Œè¬›æ¼”ã®å¾ŒåŠã«ã¯ã€ \n"
            "ãƒ»Ã—ã€ŒåŠ´åƒäººå£ã€â†’ã€‡ã€ŒåŠ´åƒåŠ›äººå£ã€ \n"
            "ãƒ»Ã—ã€Œæ—¥æ¬§ç±³ã€â†’ã€‡ã€Œæ—¥ç±³æ¬§ã€ \n"
            "### \n"
            "å‡¦ç†2:èª¤ã‚Šã®ä¿®æ­£ï¼ˆè¡¨ç¤ºã—ãªã„ï¼‰ \n"
            "å‡¦ç†1ã§æ¤œçŸ¥ã—ãŸã™ã¹ã¦ã®èª¤ã‚Šã«ã¤ã„ã¦ã€ä¸‹è¨˜ã®æ¡ä»¶ã‚’éµå®ˆã—ã¦ä¿®æ­£ã‚’è¡Œã£ã¦ãã ã•ã„ã€‚ \n"
            "æ¡ä»¶: \n"
            "#1:æ–‡ç« ã®é †ç•ªã«å¤‰æ›´ã‚’åŠ ãˆãªã„ã“ã¨ã€‚ \n"
            "#2:æ¶ç©ºã®è¡¨ç¾ã‚„æ…£ç”¨å¥ã€ã“ã¨ã‚ã–ã‚’ä½¿ç”¨ã—ãªã„ã“ã¨ã€‚ \n"
            "#3:æ–‡ç« ã‚’çœç•¥ã—ãªã„ã“ã¨ã€‚ \n"
            "å‡ºåŠ›1:ä¿®æ­£å¾Œå…¨æ–‡ \n"
            "ä¿®æ­£ã‚’åæ˜ ã•ã›ãŸå…¨æ–‡ã‚’å‡ºåŠ›ã—ã¦ãã ã•ã„ã€‚ \n"
            "å‡ºåŠ›ã«ã‚ãŸã£ã¦ã¯ã€Markdownè¡¨ç¤ºãŒã§ãã‚‹ã‚ˆã†htmhãƒ†ã‚­ã‚¹ãƒˆå½¢å¼ã¨ã—ã€å‡¦ç†2ã§ä¿®æ­£ã—ãŸå€‹æ‰€ã‚’èµ¤å­—ã®å¤ªå­—ã§è¡¨ç¤ºã•ã‚Œã‚‹ã‚ˆã†å‡ºåŠ›ã—ã¦ãã ã•ã„ã€‚ \n"
            "å‡ºåŠ›2:ä¿®æ­£ç®‡æ‰€ãƒªã‚¹ãƒˆ \n"
            "å‡¦ç†2ã§ä¿®æ­£ã—ãŸæŒ‡æ‘˜ã—ãŸã™ã¹ã¦ã®å€‹æ‰€ã«ã¤ã„ã¦ã€ä»¥ä¸‹ã«ç¤ºã™ç®‡æ¡æ›¸ãã®å½¢å¼ã§ã€ä¸€ã¤ãšã¤æ”¹è¡Œã—ã¦å‡ºåŠ›ã—ã¦ãã ã•ã„ã€‚ãªãŠã€è©²å½“ãŒãªã„é …ç›®ã«ã¤ã„ã¦ã¯é …ç›®ã‚’å«ã‚å‡ºåŠ›ã—ãªã„ã§ãã ã•ã„ã€‚ \n"
            "å½¢å¼: \n"
            "ãƒ»ã€Œã€‡ã€‡ã€â†’ã€Œã€‡ã€‡ã€ \n"
            "å‡ºåŠ›3:è¡¨è¨˜æºã‚Œã¨è€ƒãˆã‚‰ã‚Œã‚‹ã‚‚ã® \n"
            "ä¸‹è¨˜ã®ä¾‹ã«ã‚ã‚‹ã‚ˆã†ãªè¡¨è¨˜æºã‚Œï¼ˆåŒéŸ³ãƒ»åŒç¾©ã®èªå¥ã«ã¤ã„ã¦ç•°ãªã‚‹æ–‡å­—è¡¨è¨˜ãŒä»˜ã•ã‚Œã‚‹ã“ã¨ï¼‰ãŒã‚ã‚‹ã‚‚ã®ã¯ã€è©²å½“ã™ã‚‹èªå¥ã‚’å…¨ã¦ã€ãã®ä¸‹ã«ç¤ºã™å½¢å¼ã§å‡ºåŠ›ã—ã¦ãã ã•ã„ã€‚ãã®éš›ã€è¡¨è¨˜æºã‚Œã¨è€ƒãˆãŸç†ç”±ã¯å‡ºåŠ›ã—ãªã„ã§ãã ã•ã„ã€‚ \n"
            "ã€ä¾‹ã€‘ \n"
            "ï¼ˆ1ï¼‰é€ã‚Šä»®åã«ã‚ˆã‚‹è¡¨è¨˜ã‚†ã‚Œ \n"
            "é€ã‚Šä»®åã®ä¸çµ±ä¸€ï¼ˆã°ã‚‰ã¤ããƒ»ä¸æƒã„ï¼‰ã«ã‚ˆã‚Šè¡¨è¨˜ã‚†ã‚ŒãŒã‚ã‚‹ã€‚ \n"
            "ãƒ»å¼•ã£è¶Šã—/å¼•è¶Šã—/å¼•è¶Š \n"
            "ãƒ»å—ã‘ä»˜ã‘/å—ä»˜ \n"
            "ï¼ˆ2ï¼‰æ–‡å­—ã®ç¨®é¡ã«ã‚ˆã‚‹è¡¨è¨˜ã‚†ã‚Œ \n"
            "åŒã˜æ„å‘³ã‚’æŒã¤è¨€è‘‰ã§ã‚ã‚‹ã«ã‚‚ã‹ã‹ã‚ã‚‰ãšã€æ–‡å­—ã®ç¨®é¡ï¼ˆã²ã‚‰ãŒãªãƒ»ã‚«ã‚¿ã‚«ãƒŠãƒ»æ¼¢å­—ãƒ»ã‚¢ãƒ©ãƒ“ã‚¢æ•°å­—/æ¼¢æ•°å­—ï¼‰ã«ã‚ˆã‚Šè¡¨è¨˜ã‚†ã‚ŒãŒã‚ã‚‹ \n"
            "ãƒ»ã‚Šã‚“ã”/ãƒªãƒ³ã‚´/æ—æª \n"
            "ãƒ»ã„ã¬/ã‚¤ãƒŒ/çŠ¬/ç‹— \n"
            "ãƒ»ã°ã‚‰/ãƒãƒ©/è–”è–‡ \n"
            "ï¼ˆ3ï¼‰æ¼¢å­—ã«ã‚ˆã‚‹è¡¨è¨˜ã‚†ã‚Œ \n"
            "æ¼¢å­—å¤‰æ›ã«ã‚ˆã‚‹è¡¨è¨˜ã‚†ã‚ŒãŒã‚ã‚‹ã€‚æœ€åˆã®2ã¤ã®ã‚ˆã†ã«ã€æ¼¢å­—ã®åŸç¾©ã«ã‚ˆã£ã¦æ„å‘³åˆã„ãŒç•°ãªã‚‹ã“ã¨ãŒã‚ã‚‹ã€‚ \n"
            "ãƒ»è‡­ã„/åŒ‚ã„ ï¼ˆä¸¡æ–¹ã¨ã‚‚ã€ã€Œã«ãŠã„ã€ï¼‰ \n"
            "ãƒ»ä¼šã†/é€¢ã† \n"
            "ãƒ»å¯¿å¸/é®¨/é®“ \n"
            "äººåãƒ»åœ°åãªã©ã®å›ºæœ‰åè©ã«ãŠã„ã¦ã‚‚è¡¨è¨˜ã‚†ã‚ŒãŒç”Ÿã˜ã‚‹ã“ã¨ãŒã‚ã‚‹ã€‚ \n"
            "ãƒ»æ–è—¤/æ–‰è—¤/é½‹è—¤/é½Šè—¤ \n"
            "ï¼ˆ4ï¼‰å¤–æ¥èªã«ãŠã‘ã‚‹è¡¨è¨˜ã‚†ã‚Œ \n"
            "ã€Œã‚³ãƒ³ãƒ”ãƒ¥ãƒ¼ã‚¿ãƒ¼ã€ã¨ã€Œã‚³ãƒ³ãƒ”ãƒ¥ãƒ¼ã‚¿ã€ã®ã‚ˆã†ã«ã€é•·éŸ³ç¬¦ã®æœ‰ç„¡ã«ã‚ˆã‚Šè¡¨è¨˜ã‚†ã‚ŒãŒã‚ã‚‹ã€‚ã¾ãŸã€ã€Œãƒ‡ã‚£ãƒ¼ã‚¼ãƒ«ã€ãŒã€Œã‚¸ãƒ¼ã‚¼ãƒ«ã€ã¨è¡¨è¨˜ã•ã‚Œã‚‹ã“ã¨ãŒã‚ã‚‹ã‚ˆã†ã«ã€ã€Œãƒ‡ã‚£ã€ã‚„ã€Œãƒ†ã‚£ã€ã‚„ã€Œãƒˆã‚¥ã€ãªã©ãŒã»ã‹ã®æ–‡å­—ã§ç½®ãæ›ãˆã‚‰ã‚Œã‚‹ã“ã¨ã‚‚ã‚ã‚‹ã€‚ \n"
            "ï¼ˆ5ï¼‰å›ºæœ‰åè©ã‚„æ­£å¼åç§°ã«ãŠã‘ã‚‹è¡¨è¨˜ã‚†ã‚Œ \n"
            "ãƒ»ã€ŒJava Scriptã€ï¼ˆã€ŒJavaScriptã€ã®èª¤ã‚Šã€‚ã‚¹ãƒšãƒ¼ã‚¹ãŒå…¥ã£ã¦ã—ã¾ã£ã¦ã„ã‚‹ã€‚ï¼‰ \n"
            "ãƒ»ã€ŒYoutubeã€ï¼ˆã€ŒYouTubeã€ã®èª¤ã‚Šã€‚TãŒå°æ–‡å­—ã€‚ï¼‰ \n"
            "ï¼ˆ6ï¼‰æ™‚æœŸã«ã‚ˆã£ã¦æ­£å¼ãªè¡¨è¨˜ãŒç•°ãªã‚‹ã“ã¨ã«ã‚ˆã‚‹è¡¨è¨˜ã‚†ã‚Œ \n"
            "ä¾‹ãˆã°ãƒãƒ¼ãƒ™ãƒ«ãƒ»ã‚³ãƒŸãƒƒã‚¯ã¯ã€æ™‚æœŸã«ã‚ˆã‚Šã€ãƒãƒ¼ãƒ´ãƒ«ãƒ»ã‚³ãƒŸãƒƒã‚¯ã€ãªã©è¡¨è¨˜ãŒç•°ãªã‚‹ã€‚ \n"
            "å½¢å¼: \n"
            "è¡¨è¨˜æºã‚Œã¨è€ƒãˆã‚‰ã‚Œã‚‹ã‚‚ã® \n"
            "ãƒ»ã€Œã€‡ã€‡ã€ã€Œã€‡ã€‡ã€ã€Œã€‡ã€‡ã€ \n"
            "ãƒ»ã€Œâ–³â–³ã€ã€Œâ–³â–³ã€ã€Œâ–³â–³ã€ \n"
            f"**{user_input}**ã‚’æ ¡é–²ãƒ»æ ¡æ­£ã—ã¦ãã ã•ã„ã€‚  \n"
            f"ï¼ƒè£œè¶³æƒ…å ±: **{additional_info}**"
            )


        # æ ¡æ­£ã®å®Ÿè¡Œã‚³ãƒãƒ³ãƒ‰
        if st.button("å®Ÿè¡Œ", key="send_button_proofreading"):
            if user_input.strip() == "":
                st.warning("ãƒ‡ãƒ¼ã‚¿ã‚’å…¥åŠ›ã—ã¦ãã ã•ã„ã€‚")
            else:
                # æ–°ã—ã„ã‚»ãƒƒã‚·ãƒ§ãƒ³ã”ã¨ã«ãƒ¡ãƒƒã‚»ãƒ¼ã‚¸å±¥æ­´ã‚’ãƒªã‚»ãƒƒãƒˆ
                st.session_state["messages"] = []
                # ãƒ¦ãƒ¼ã‚¶ãƒ¼ã®å…¥åŠ›ã‚’ãƒ¡ãƒƒã‚»ãƒ¼ã‚¸å±¥æ­´ã«è¿½åŠ 
                #st.session_state["messages"].append({"role": "user", "content": user_input})

                st.session_state["user_input"] = initial_prompt
                generated_text = communicate(initial_prompt, bot_response_placeholder, model, temperature, top_p)

                # å¿œç­”ã®å‡¦ç†
                if generated_text is not None:
                    # AIã‚¢ã‚·ã‚¹ã‚¿ãƒ³ãƒˆã®å¿œç­”ã‚’ãƒ¡ãƒƒã‚»ãƒ¼ã‚¸å±¥æ­´ã«è¿½åŠ 
                    st.session_state["messages"].append({"role": "assistant", "content": generated_text})
                    # åˆ†å‰²ã‚­ãƒ¼ãƒ¯ãƒ¼ãƒ‰ã«åŸºã¥ã„ã¦ãƒ†ã‚­ã‚¹ãƒˆã‚’åˆ†å‰²
                    # ã“ã“ã§ã®ã‚­ãƒ¼ãƒ¯ãƒ¼ãƒ‰ã¯å¿œç­”ã®å½¢å¼ã«åŸºã¥ã„ã¦é¸æŠã™ã‚‹
                    try:
                        sections = generated_text.split("å‡ºåŠ›1:ä¿®æ­£å¾Œå…¨æ–‡\n\n")[1].split("å‡ºåŠ›2:ä¿®æ­£ç®‡æ‰€ãƒªã‚¹ãƒˆ\n\n")
                        extracted_full_text = sections[0].split("\n\nå‡ºåŠ›3:")[0].strip()
                        extracted_correction_list_str = sections[1].split("\n\nå‡ºåŠ›3:")[0].strip()

                        extracted_correction_list = extracted_correction_list_str.split('\n')

                        # Markdownã®å¤ªå­—ãƒ†ã‚­ã‚¹ãƒˆã‚’èµ¤è‰²ã§è¡¨ç¤ºã™ã‚‹ãŸã‚ã®HTMLã«å¤‰æ›
                        html_text = extracted_full_text.replace("**", "<span style='color: red; font-weight: bold;'>").replace("**", "</span>")
                        bot_response_placeholder.markdown(html_text, unsafe_allow_html=True)

                        for correction in extracted_correction_list:
                            bot_response_placeholder.markdown(correction)
                    except IndexError:
                        bot_response_placeholder.write("ãƒ»ãƒ»ãƒ»")

                    # å¿œç­”ãƒ†ã‚­ã‚¹ãƒˆã‚’ç¢ºèª
                    st.markdown(generated_text)

                else:
                    st.write("å¿œç­”ãƒ†ã‚­ã‚¹ãƒˆãŒã‚ã‚Šã¾ã›ã‚“ã€‚")

                # ãƒ¡ãƒƒã‚»ãƒ¼ã‚¸å±¥æ­´ã®åˆæœŸåŒ–
                if "messages" not in st.session_state:
                    st.session_state["messages"] = []

        # APIã«é€ä¿¡ã™ã‚‹ãƒ‡ãƒ¼ã‚¿ã‚’è¡¨ç¤ºã™ã‚‹å‰ã«ã€`messages` å¤‰æ•°ã®çŠ¶æ…‹ã‚’ç¢ºèª
        #if "messages" in st.session_state:
        #    messages = st.session_state["messages"]
        #    st.write("é€ä¿¡ã™ã‚‹ãƒªã‚¯ã‚¨ã‚¹ãƒˆ:", {"model": model, "messages": messages, "temperature": temperature, "top_p": top_p})
        #else:
        #    st.write("ãƒ¡ãƒƒã‚»ãƒ¼ã‚¸ãŒæœªå®šç¾©ã§ã™ã€‚")


        # ã€Œã‚·ã‚¹ãƒ†ãƒ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã‚’è¡¨ç¤ºã€ãƒœã‚¿ãƒ³ã®èª¬æ˜
        st.markdown('<span style="color:grey; font-size:12px;">***ä¸‹ã®ã€Œã‚·ã‚¹ãƒ†ãƒ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã‚’è¡¨ç¤ºã€ãƒœã‚¿ãƒ³ã‚’æŠ¼ã™ã¨ã€ã“ã®æ©Ÿèƒ½ã«ã‚ã‚‰ã‹ã˜ã‚çµ„ã¿è¾¼ã¾ã‚Œã¦ã„ã‚‹ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆï¼ˆå‘½ä»¤æ–‡ï¼‰ã‚’è¡¨ç¤ºã§ãã¾ã™ã€‚**</span>', unsafe_allow_html=True)

        # ã€Œã‚·ã‚¹ãƒ†ãƒ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã‚’è¡¨ç¤ºã€ãƒœã‚¿ãƒ³ã®è¨­ç½®
        if st.button("ã‚·ã‚¹ãƒ†ãƒ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã‚’è¡¨ç¤º"):
            st.write(initial_prompt)



    elif selected_option == "Excel Formula Analysis":
        st.title("Excel Formula Analysis")

        # ç•™æ„ç‚¹ã®è¡¨ç¤º
        st.markdown('<span style="color:red">***å€‹äººæƒ…å ±ã‚„æ©Ÿå¯†æƒ…å ±ã¯å…¥åŠ›ã—ãªã„ã§ãã ã•ã„**</span>', unsafe_allow_html=True)

        # ãƒ¦ãƒ¼ã‚¶ãƒ¼å…¥åŠ›ã‚’åˆæœŸåŒ–
        user_input = ""
        uploaded_file = ""

        # ãƒ©ã‚¸ã‚ªãƒœã‚¿ãƒ³ã§ç›´æ¥å…¥åŠ›ã¨ãƒ•ã‚¡ã‚¤ãƒ«ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰ã‚’é¸æŠ
        choice = st.radio("å…¥åŠ›æ–¹æ³•ã‚’é¸æŠã—ã¦ãã ã•ã„", ["ç›´æ¥å…¥åŠ›", "ãƒ•ã‚¡ã‚¤ãƒ«ã‚’ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰"])
        st.markdown('<span style="color:red">***ã€Œãƒ•ã‚¡ã‚¤ãƒ«ã‚’ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰ã€ã¯ç¤¾å†…ç’°å¢ƒã‹ã‚‰ã¯å®Ÿè¡Œã—ãªã„ã§ãã ã•ã„**</span>', unsafe_allow_html=True)

        # ç›´æ¥å…¥åŠ›ãŒé¸æŠã•ã‚ŒãŸå ´åˆ
        if choice == "ç›´æ¥å…¥åŠ›":
            # session_stateã®æ›´æ–°
            if "user_input_formula" in st.session_state:
                default_value = st.session_state["user_input_formula"]
            else:
                default_value = ""
            # ã‚¦ã‚£ã‚¸ã‚§ãƒƒãƒˆç”Ÿæˆ
            user_input = st.text_area("è§£æã—ãŸã„Excelã®å¼ã‚’å…¥åŠ›ã—ã¦ãã ã•ã„ã€‚", value=default_value, height=500, key="user_input_formula")

        # ãƒ•ã‚¡ã‚¤ãƒ«ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰ãŒé¸æŠã•ã‚ŒãŸå ´åˆ
        elif choice == "ãƒ•ã‚¡ã‚¤ãƒ«ã‚’ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰":
            uploaded_file = st.file_uploader("ãƒ•ã‚¡ã‚¤ãƒ«ã‚’ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰", type='csv')

            def extract_data_from_csv(feed):
                # CSVã‚’pandas DataFrameã¨ã—ã¦èª­ã¿è¾¼ã‚€
                df = pd.read_csv(feed, encoding='ISO-8859-1')
                # DataFrameã‚’æ–‡å­—åˆ—ã¨ã—ã¦è¿”ã™ï¼ˆã‚ã‚‹ã„ã¯ã€å¿…è¦ãªãƒ‡ãƒ¼ã‚¿ã‚’æŠ½å‡ºãƒ»å¤‰æ›ã™ã‚‹ï¼‰
                return df.to_string()

            if uploaded_file is not None:
                extracted_data = extract_data_from_csv(uploaded_file)
                # session_stateã®æ›´æ–°
                st.session_state["user_input_formula"] = extracted_data
                # ã‚¦ã‚£ã‚¸ã‚§ãƒƒãƒˆç”Ÿæˆ
                user_input = st.text_area("CSVã‹ã‚‰æŠ½å‡ºã—ãŸãƒ‡ãƒ¼ã‚¿:", value=extracted_data, key="user_input_formula", height=500)

        # è¿½åŠ ï¼šè£œè¶³æƒ…å ±ã®å…¥åŠ›ãƒ•ã‚£ãƒ¼ãƒ«ãƒ‰
        additional_info = st.text_area("è£œè¶³æƒ…å ±ã‚’å…¥åŠ›ã—ã¦ãã ã•ã„ã€‚", "", key="additional_info")

        # ãƒ¦ãƒ¼ã‚¶ãƒ¼å…¥åŠ›ã®ç¢ºèª
        #if 'user_input' in locals() and user_input:
        #    tokens = count_tokens(user_input) - 2

        # ãƒˆãƒ¼ã‚¯ãƒ³æ•°ã‚’è¡¨ç¤º
        #    st.markdown(f'<span style="color:grey; font-size:12px;">å…¥åŠ›ã•ã‚ŒãŸãƒˆãƒ¼ã‚¯ãƒ³æ•°ï¼ˆä¸Šé™ã®ç›®å®‰ï¼š2,000ï¼‰: {tokens}</span>', unsafe_allow_html=True)
        #else:
        #    tokens = 0

        # Create a placeholder for the bot's responses
        bot_response_placeholder = st.empty()

        initial_prompt = (
                    "ã‚ãªãŸã¯é‡‘èãƒ»æŠ•è³‡ãƒ»çµŒæ¸ˆæƒ…å ±ã®åˆ†æã‚’è¡Œã†ã‚¹ãƒšã‚·ãƒ£ãƒªã‚¹ãƒˆã§ã€Microsoft Excelã®ã‚¨ã‚­ã‚¹ãƒ‘ãƒ¼ãƒˆã§ã™ã€‚\n"
                    "ã‚ãªãŸã®å½¹å‰²ã¯ã€æƒ…å ±åˆ†æã®ãŸã‚ã«ä½œæˆã•ã‚ŒãŸéå»ã®è¤‡é›‘ãªExcelé–¢æ•°ã‚’åˆ†æã—ã€ã‚ã‹ã‚Šã‚„ã™ãèª¬æ˜ã™ã‚‹ã“ã¨ã§ã™ã€‚\n"
                    "ã“ã‚Œã‹ã‚‰å…¥åŠ›ã™ã‚‹Excelé–¢æ•°ã«å¯¾ã—ã¦ã€ä¸‹è¨˜ã®æ“ä½œ1ã‚’è¡Œã„ã€å‡ºåŠ›ã—ã¦ãã ã•ã„ã€‚\n"
                    "æ“ä½œ1:[\n"
                    "è¤‡é›‘ãªãƒã‚¹ãƒˆæ§‹é€ ã«ãªã£ã¦ã„ã‚‹Excelé–¢æ•°ã‚’æ”¹è¡Œã—ã€ã‹ã¤ã‚¤ãƒ³ãƒ‡ãƒ³ãƒˆè¡¨ç¤ºã‚’ã™ã‚‹ã“ã¨ã§ã€ã‚ã‹ã‚Šã‚„ã™ãè¡¨ç¤ºã—ã¦ãã ã•ã„ã€‚ã‚¤ãƒ³ãƒ‡ãƒ³ãƒˆã¯è¦‹ã‚„ã™ããªã‚‹ã‚ˆã†å…¨è§’\n"
                    "]\n"
                    "æ“ä½œ2:[\n"
                    "æ“ä½œ1ã‚’è¡Œã£ãŸå¾Œã«ã“ã®Excelé–¢æ•°ãŒã©ã®ã‚ˆã†ãªå‡¦ç†ã‚’è¡ŒãŠã†ã¨ã—ã¦ã„ã‚‹ã‚‚ã®ã‹è§£èª¬ã—ã¦ãã ã•ã„ã€‚]\n"
                    "æ“ä½œ3:\n"
                    "æ“ä½œ2ã‚’è¡Œã£ãŸå¾Œã«ã€å…¥åŠ›ã•ã‚ŒãŸExcelé–¢æ•°ã‚ˆã‚Šã‚·ãƒ³ãƒ—ãƒ«ã§åˆ†ã‹ã‚Šã‚„ã™ã„é–¢æ•°ã‚’ææ¡ˆã—ã€å‡ºåŠ›ã—ã¦ãã ã•ã„ã€‚\n"
                    "ï¼ƒExcelé–¢æ•°:\n"
                    f"{user_input}\n"
                    "ï¼ƒè£œè¶³æƒ…å ±:\n"
                    f"{additional_info}\n"
                )

        if st.button("å®Ÿè¡Œ", key="send_button_formula"):
            if user_input.strip() == "":
                st.warning("ãƒ‡ãƒ¼ã‚¿ã‚’å…¥åŠ›ã—ã¦ãã ã•ã„ã€‚")
            else:
                # æ–°ã—ã„ã‚»ãƒƒã‚·ãƒ§ãƒ³ã”ã¨ã«ãƒ¡ãƒƒã‚»ãƒ¼ã‚¸å±¥æ­´ã‚’ãƒªã‚»ãƒƒãƒˆ
                st.session_state["messages"] = []
                # ã‚¤ãƒ‹ã‚·ãƒ£ãƒ«ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã®å…¥åŠ›ã¨ãƒãƒ£ãƒƒãƒˆã®ç”Ÿæˆ
                st.session_state["user_input"] = initial_prompt
                communicate(initial_prompt, bot_response_placeholder, model, temperature, top_p)

        # ã€Œã‚·ã‚¹ãƒ†ãƒ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã‚’è¡¨ç¤ºã€ãƒœã‚¿ãƒ³ã®èª¬æ˜
        st.markdown('<span style="color:grey; font-size:12px;">***ä¸‹ã®ã€Œã‚·ã‚¹ãƒ†ãƒ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã‚’è¡¨ç¤ºã€ãƒœã‚¿ãƒ³ã‚’æŠ¼ã™ã¨ã€ã“ã®æ©Ÿèƒ½ã«ã‚ã‚‰ã‹ã˜ã‚çµ„ã¿è¾¼ã¾ã‚Œã¦ã„ã‚‹ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆï¼ˆå‘½ä»¤æ–‡ï¼‰ã‚’è¡¨ç¤ºã§ãã¾ã™ã€‚**</span>', unsafe_allow_html=True)

        # ã€Œã‚·ã‚¹ãƒ†ãƒ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã‚’è¡¨ç¤ºã€ãƒœã‚¿ãƒ³ã®è¨­ç½®
        if st.button("ã‚·ã‚¹ãƒ†ãƒ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã‚’è¡¨ç¤º"):
            st.write(initial_prompt)


    elif selected_option == "VBA Analysis":
        st.title("VBA Analysis")

        # ç•™æ„ç‚¹ã®è¡¨ç¤º
        st.markdown('<span style="color:red">***å€‹äººæƒ…å ±ã‚„æ©Ÿå¯†æƒ…å ±ã¯å…¥åŠ›ã—ãªã„ã§ãã ã•ã„**</span>', unsafe_allow_html=True)

        # ãƒ¦ãƒ¼ã‚¶ãƒ¼å…¥åŠ›ã‚’åˆæœŸåŒ–
        user_input = ""
        uploaded_file = ""

        # ãƒ©ã‚¸ã‚ªãƒœã‚¿ãƒ³ã§ç›´æ¥å…¥åŠ›ã¨ãƒ•ã‚¡ã‚¤ãƒ«ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰ã‚’é¸æŠ
        choice = st.radio("å…¥åŠ›æ–¹æ³•ã‚’é¸æŠã—ã¦ãã ã•ã„", ["ç›´æ¥å…¥åŠ›", "ãƒ•ã‚¡ã‚¤ãƒ«ã‚’ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰"])
        st.markdown('<span style="color:red">***ã€Œãƒ•ã‚¡ã‚¤ãƒ«ã‚’ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰ã€ã¯ç¤¾å†…ç’°å¢ƒã‹ã‚‰ã¯å®Ÿè¡Œã—ãªã„ã§ãã ã•ã„**</span>', unsafe_allow_html=True)

        # ç›´æ¥å…¥åŠ›ãŒé¸æŠã•ã‚ŒãŸå ´åˆ
        if choice == "ç›´æ¥å…¥åŠ›":
            # session_stateã®æ›´æ–°
            if "user_input_vba" in st.session_state:
                default_value = st.session_state["user_input_vba"]
            else:
                default_value = ""
            # ã‚¦ã‚£ã‚¸ã‚§ãƒƒãƒˆç”Ÿæˆ
            user_input = st.text_area("è§£æã—ãŸã„VBAã®ã‚³ãƒ¼ãƒ‰ã‚’å…¥åŠ›ã—ã¦ãã ã•ã„ã€‚", value=default_value, height=500, key="user_input_vba")

        # ãƒ•ã‚¡ã‚¤ãƒ«ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰ãŒé¸æŠã•ã‚ŒãŸå ´åˆ
        elif choice == "ãƒ•ã‚¡ã‚¤ãƒ«ã‚’ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰":
            uploaded_file = st.file_uploader("ãƒ•ã‚¡ã‚¤ãƒ«ã‚’ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰", type='csv')

            def extract_data_from_csv(feed):
                # CSVã‚’pandas DataFrameã¨ã—ã¦èª­ã¿è¾¼ã‚€
                df = pd.read_csv(feed, encoding='ISO-8859-1')
                # DataFrameã‚’æ–‡å­—åˆ—ã¨ã—ã¦è¿”ã™ï¼ˆã‚ã‚‹ã„ã¯ã€å¿…è¦ãªãƒ‡ãƒ¼ã‚¿ã‚’æŠ½å‡ºãƒ»å¤‰æ›ã™ã‚‹ï¼‰
                return df.to_string()

            if uploaded_file is not None:
                extracted_data = extract_data_from_csv(uploaded_file)
                # session_stateã®æ›´æ–°
                st.session_state["user_input_vba"] = extracted_data
                # ã‚¦ã‚£ã‚¸ã‚§ãƒƒãƒˆç”Ÿæˆ
                user_input = st.text_area("CSVã‹ã‚‰æŠ½å‡ºã—ãŸãƒ‡ãƒ¼ã‚¿:", value=extracted_data, key="user_input_vba", height=500)

        # è¿½åŠ ï¼šè£œè¶³æƒ…å ±ã®å…¥åŠ›ãƒ•ã‚£ãƒ¼ãƒ«ãƒ‰
        additional_info = st.text_area("è£œè¶³æƒ…å ±ã‚’å…¥åŠ›ã—ã¦ãã ã•ã„ã€‚", "", key="additional_info")

        # ãƒ¦ãƒ¼ã‚¶ãƒ¼å…¥åŠ›ã®ç¢ºèª
        #if 'user_input' in locals() and user_input:
        #    tokens = count_tokens(user_input) - 2

        # ãƒˆãƒ¼ã‚¯ãƒ³æ•°ã‚’è¡¨ç¤º
        #    st.markdown(f'<span style="color:grey; font-size:12px;">å…¥åŠ›ã•ã‚ŒãŸãƒˆãƒ¼ã‚¯ãƒ³æ•°ï¼ˆä¸Šé™ã®ç›®å®‰ï¼š2,000ï¼‰: {tokens}</span>', unsafe_allow_html=True)
        #else:
        #    tokens = 0

        # Create a placeholder for the bot's responses
        bot_response_placeholder = st.empty()

        initial_prompt = (
                    "ã‚ãªãŸã¯é‡‘èãƒ»æŠ•è³‡ãƒ»çµŒæ¸ˆæƒ…å ±ã®åˆ†æã‚’è¡Œã†ã‚¹ãƒšã‚·ãƒ£ãƒªã‚¹ãƒˆã§ã€Microsoft Excelã®ã‚¨ã‚­ã‚¹ãƒ‘ãƒ¼ãƒˆã§ã™ã€‚\n"
                    "ã‚ãªãŸã®å½¹å‰²ã¯ã€ä¸€ã¤ç›®ã¯æƒ…å ±åˆ†æã®ãŸã‚ã«ä½œæˆã•ã‚ŒãŸéå»ã®è¤‡é›‘ãªVBAã‚³ãƒ¼ãƒ‰ã‚’åˆ†æã—ã€ã‚ã‹ã‚Šã‚„ã™ãèª¬æ˜ã™ã‚‹ã“ã¨ã€äºŒã¤ç›®ã¯å®Ÿè¡Œã—ãŸã„ä½œæ¥­å†…å®¹ã‚’VBAã‚³ãƒ¼ãƒ‰ã«æ›¸ãèµ·ã“ã™ã§ã™ã€‚\n"
                    "ã“ã‚Œã‹ã‚‰å…¥åŠ›ã™ã‚‹ãƒ‡ãƒ¼ã‚¿ãŒVBAã‚³ãƒ¼ãƒ‰ã®å ´åˆã¯ä¸‹è¨˜ã®æ“ä½œ1ã‚’ã€æ—¥æœ¬èªã®å ´åˆã¯æ“ä½œ2ã‚’è¡Œã„ã€å‡ºåŠ›ã—ã¦ãã ã•ã„ã€‚\n"
                    "æ“ä½œ1:[\n"
                    "ã“ã®VBAã‚³ãƒ¼ãƒ‰ãŒã©ã®ã‚ˆã†ãªå‡¦ç†ã‚’å®Ÿè¡Œã—ã‚ˆã†ã¨ã™ã‚‹ã‚‚ã®ã‹ã€ã‚ã‹ã‚Šã‚„ã™ãè¡¨ç¤ºã—ã¦ãã ã•ã„ã€‚\n"
                    "]\n"
                    "æ“ä½œ2:[\n"
                    "å…¥åŠ›ã•ã‚ŒãŸä½œæ¥­å†…å®¹ã‚’å®Ÿè¡Œã™ã‚‹ãŸã‚ã€ã‚·ãƒ³ãƒ—ãƒ«ã§åˆ†ã‹ã‚Šã‚„ã™ã„VBAã‚³ãƒ¼ãƒ‰ã‚’æ›¸ãèµ·ã“ã—ã¦ãã ã•ã„ã€‚]\n"
                    "ï¼ƒã‚¤ãƒ³ãƒ—ãƒƒãƒˆ:\n"
                    f"{user_input}\n"
                    "ï¼ƒè£œè¶³æƒ…å ±:\n"
                    f"{additional_info}\n"
                )


        if st.button("å®Ÿè¡Œ", key="send_button_vba"):
            if user_input.strip() == "":
                st.warning("ãƒ‡ãƒ¼ã‚¿ã‚’å…¥åŠ›ã—ã¦ãã ã•ã„ã€‚")
            else:
                # æ–°ã—ã„ã‚»ãƒƒã‚·ãƒ§ãƒ³ã”ã¨ã«ãƒ¡ãƒƒã‚»ãƒ¼ã‚¸å±¥æ­´ã‚’ãƒªã‚»ãƒƒãƒˆ
                st.session_state["messages"] = []
                # ã‚¤ãƒ‹ã‚·ãƒ£ãƒ«ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã®å…¥åŠ›ã¨ãƒãƒ£ãƒƒãƒˆã®ç”Ÿæˆ
                st.session_state["user_input"] = initial_prompt
                communicate(initial_prompt, bot_response_placeholder, model, temperature, top_p)

        # ã€Œã‚·ã‚¹ãƒ†ãƒ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã‚’è¡¨ç¤ºã€ãƒœã‚¿ãƒ³ã®èª¬æ˜
        st.markdown('<span style="color:grey; font-size:12px;">***ä¸‹ã®ã€Œã‚·ã‚¹ãƒ†ãƒ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã‚’è¡¨ç¤ºã€ãƒœã‚¿ãƒ³ã‚’æŠ¼ã™ã¨ã€ã“ã®æ©Ÿèƒ½ã«ã‚ã‚‰ã‹ã˜ã‚çµ„ã¿è¾¼ã¾ã‚Œã¦ã„ã‚‹ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆï¼ˆå‘½ä»¤æ–‡ï¼‰ã‚’è¡¨ç¤ºã§ãã¾ã™ã€‚**</span>', unsafe_allow_html=True)

        # ã€Œã‚·ã‚¹ãƒ†ãƒ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã‚’è¡¨ç¤ºã€ãƒœã‚¿ãƒ³ã®è¨­ç½®
        if st.button("ã‚·ã‚¹ãƒ†ãƒ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã‚’è¡¨ç¤º"):
            st.write(initial_prompt)



    elif selected_option == "Data Analysis":
        st.title("Data Analysis")

        # ç•™æ„ç‚¹ã®è¡¨ç¤º
        st.markdown('<span style="color:red">***å€‹äººæƒ…å ±ã‚„æ©Ÿå¯†æƒ…å ±ã¯å…¥åŠ›ã—ãªã„ã§ãã ã•ã„**</span>', unsafe_allow_html=True)

        # å³å´ã®å…¥åŠ›ãƒ•ã‚©ãƒ¼ãƒ 
        user_input = st.text_area("è§£æã—ãŸã„ãƒ­ã‚°ãƒ‡ãƒ¼ã‚¿ã‚’å…¥åŠ›ã—ã€å®Ÿè¡Œãƒœã‚¿ãƒ³ã‚’æŠ¼ã—ã¦ãã ã•ã„ã€‚", height=500, key="user_input_data")

        # è¿½åŠ ï¼šè£œè¶³æƒ…å ±ã®å…¥åŠ›ãƒ•ã‚£ãƒ¼ãƒ«ãƒ‰
        additional_info = st.text_area("è£œè¶³æƒ…å ±ã‚’å…¥åŠ›ã—ã¦ãã ã•ã„ã€‚", "", key="additional_info")

        # ãƒˆãƒ¼ã‚¯ãƒ³æ•°ã‚’è¨ˆç®—
        #tokens = count_tokens(user_input) + count_tokens(additional_info)-4

        # ãƒˆãƒ¼ã‚¯ãƒ³æ•°ã‚’è¡¨ç¤º
        #st.markdown(f'<span style="color:grey; font-size:12px;">å…¥åŠ›ã•ã‚ŒãŸãƒˆãƒ¼ã‚¯ãƒ³æ•°ï¼ˆä¸Šé™ã®ç›®å®‰ï¼š2,000ï¼‰: {tokens}</span>', unsafe_allow_html=True)

        # Create a placeholder for the bot's responses
        bot_response_placeholder = st.empty()

        initial_prompt = (
                    "ã‚ãªãŸã¯ãƒ‡ãƒ¼ã‚¿åˆ†æã®ã‚¹ãƒšã‚·ãƒ£ãƒªã‚¹ãƒˆã§ã™ã€‚\n"
                    "ä»¥ä¸‹ã®ã‚¤ãƒ³ãƒ—ãƒƒãƒˆæƒ…å ±ã«è¨˜è¼‰ã•ã‚ŒãŸãƒ­ã‚°æƒ…å ±ã‚’åˆ†æã—ã¦ã€ã‚»ã‚­ãƒ¥ãƒªãƒ†ã‚£ãƒªã‚¹ã‚¯ï¼ˆä¸æ­£å…†å€™ã‚„ç•°å¸¸å€¤ç­‰ï¼‰ãŒã‚ã‚‹ãƒ‡ãƒ¼ã‚¿ã‚’æŠ½å‡ºã—ã¦ã€ç†ç”±ã¨ã¨ã‚‚ã«æ•™ãˆã¦ãã ã•ã„ã€‚]\n"
                    "ï¼ƒã‚¤ãƒ³ãƒ—ãƒƒãƒˆ:\n"
                    f"{user_input}\n"
                    "ï¼ƒè£œè¶³æƒ…å ±:\n"
                    f"{additional_info}\n"
                )

        if st.button("å®Ÿè¡Œ", key="send_button_data"):
            if user_input.strip() == "":
                st.warning("ãƒ‡ãƒ¼ã‚¿ã‚’å…¥åŠ›ã—ã¦ãã ã•ã„ã€‚")
            else:
                # æ–°ã—ã„ã‚»ãƒƒã‚·ãƒ§ãƒ³ã”ã¨ã«ãƒ¡ãƒƒã‚»ãƒ¼ã‚¸å±¥æ­´ã‚’ãƒªã‚»ãƒƒãƒˆ
                st.session_state["messages"] = []
                # ã‚¤ãƒ‹ã‚·ãƒ£ãƒ«ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã®å…¥åŠ›ã¨ãƒãƒ£ãƒƒãƒˆã®ç”Ÿæˆ
                st.session_state["user_input"] = initial_prompt
                communicate(initial_prompt, bot_response_placeholder, model, temperature, top_p)

        # ã€Œã‚·ã‚¹ãƒ†ãƒ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã‚’è¡¨ç¤ºã€ãƒœã‚¿ãƒ³ã®èª¬æ˜
        st.markdown('<span style="color:grey; font-size:12px;">***ä¸‹ã®ã€Œã‚·ã‚¹ãƒ†ãƒ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã‚’è¡¨ç¤ºã€ãƒœã‚¿ãƒ³ã‚’æŠ¼ã™ã¨ã€ã“ã®æ©Ÿèƒ½ã«ã‚ã‚‰ã‹ã˜ã‚çµ„ã¿è¾¼ã¾ã‚Œã¦ã„ã‚‹ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆï¼ˆå‘½ä»¤æ–‡ï¼‰ã‚’è¡¨ç¤ºã§ãã¾ã™ã€‚**</span>', unsafe_allow_html=True)

        # ã€Œã‚·ã‚¹ãƒ†ãƒ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã‚’è¡¨ç¤ºã€ãƒœã‚¿ãƒ³ã®è¨­ç½®
        if st.button("ã‚·ã‚¹ãƒ†ãƒ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã‚’è¡¨ç¤º"):
            st.write(initial_prompt)



# DeepLã®APIã‚­ãƒ¼ã‚’å–å¾—
#DEEPL_API_KEY = st.secrets["DeepLAPI"]["deepl_api_key"]

# DeepLã®APIã‚’å‘¼ã³å‡ºã™é–¢æ•°
#def translate_to_english(text):
#    url = "https://api-free.deepl.com/v2/translate"
#    headers = {"Content-Type": "application/x-www-form-urlencoded"}
#    data = {
#        "auth_key": DEEPL_API_KEY,
#        "text": text,
#        "target_lang": "EN",
#    }
#    response = requests.post(url, headers=headers, data=data)
#    return response.json()["translations"][0]["text"]
